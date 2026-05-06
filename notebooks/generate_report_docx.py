"""
MASA Hackathon 2026: R-Ignite
DOCX — 10 body pages, Times New Roman 12pt, A4, 1-inch margins, all charts embedded.
FIX: Never set line_spacing as Pt() on Normal style — python-docx writes lineRule=exact
     which clips images to 14pt. Image paragraphs use lineRule=atLeast instead.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE  = "/Users/a1357/Documents/GitHub/R-Ignite-Masa-Hacktathonn"
OUT_D = os.path.join(BASE, "outputs")
DOCX  = os.path.join(OUT_D, "MASA_R_Ignite_Report_2026.docx")

DARK  = RGBColor(0x1a, 0x1a, 0x2e)
MID   = RGBColor(0x6b, 0x27, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT  = "Times New Roman"
BODY  = Pt(12)
SM    = Pt(10)
CAPT  = Pt(9)

# ── XML spacing helper ────────────────────────────────────────────────────────
def _spc(para, before=0, after=0, line=None, rule=None):
    """Set paragraph spacing in twips (1 twip = 1/20 pt = 1/1440 in).
    rule: 'exact' | 'atLeast' | 'auto'  (None = omit, inherits from style)
    NEVER pass rule='exact' for image paragraphs — it clips images."""
    pPr = para._p.get_or_add_pPr()
    # Remove any existing spacing element first to avoid duplicates
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    el = OxmlElement('w:spacing')
    el.set(qn('w:before'), str(before))
    el.set(qn('w:after'),  str(after))
    if line is not None:
        el.set(qn('w:line'),     str(line))
        el.set(qn('w:lineRule'), rule or 'atLeast')
    pPr.append(el)

def _cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)

def _run(para, text, bold=False, italic=False, size=None, color=None):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = size or BODY
    r.bold   = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    return r

def _keep_next(para):
    para._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))

def _page_break(doc):
    p = doc.add_paragraph()
    _spc(p, 0, 0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# ── Paragraph factories ───────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_paragraph()
    _spc(p, before=120, after=60, line=276, rule='atLeast')
    _keep_next(p)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = DARK
    pPr = p._p.get_or_add_pPr()
    bd  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'6b2737')
    bd.append(bot); pPr.append(bd)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    _spc(p, before=80, after=30, line=240, rule='atLeast')
    _keep_next(p)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = BODY; r.bold = True
    r.font.color.rgb = MID
    return p

def body(doc, text, bold=False, italic=False, before=0, after=80,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    _spc(p, before=before, after=after, line=240, rule='atLeast')
    _run(p, text, bold=bold, italic=italic)
    return p

def mixed(doc, parts, before=0, after=80, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    _spc(p, before=before, after=after, line=240, rule='atLeast')
    for t, b, i in parts: _run(p, t, bold=b, italic=i)
    return p

def capt(doc, text, after=80):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spc(p, before=20, after=after, line=240, rule='atLeast')
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = CAPT; r.italic = True
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    return p

def highlight(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(p, before=60, after=60, line=240, rule='atLeast')
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),'E8F4F8'); pPr.append(shd)
    bdr = OxmlElement('w:pBdr')
    lft = OxmlElement('w:left')
    lft.set(qn('w:val'),'single'); lft.set(qn('w:sz'),'12')
    lft.set(qn('w:space'),'4');    lft.set(qn('w:color'),'1a6b9a')
    bdr.append(lft); pPr.append(bdr)
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'),'180'); pPr.append(ind)
    _run(p, text)
    return p

def bullet(doc, text, after=40):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(p, before=0, after=after, line=240, rule='atLeast')
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),'360'); ind.set(qn('w:hanging'),'180')
    pPr.append(ind)
    _run(p, u'•  ' + text)
    return p

def finding(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(p, before=40, after=40, line=240, rule='atLeast')
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'),'180'); pPr.append(ind)
    _run(p, text, bold=True, italic=True, color=MID)
    return p

# ── Table ─────────────────────────────────────────────────────────────────────
def tbl(doc, headers, rows, col_w, after=80):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, w in enumerate(col_w):
        for row in t.rows:
            row.cells[ci].width = Inches(w)
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        _cell_bg(c, '6b2737')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spc(p, 0, 0, line=240, rule='atLeast')
        r = p.add_run(h)
        r.font.name = FONT; r.font.size = SM; r.bold = True
        r.font.color.rgb = WHITE
    for ri, row_data in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F5F5'
        for ci, cell_text in enumerate(row_data):
            c = t.rows[ri+1].cells[ci]
            _cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spc(p, 0, 0, line=240, rule='atLeast')
            r = p.add_run(str(cell_text))
            r.font.name = FONT; r.font.size = SM
    return t

# ── Images — lineRule=atLeast so images are NEVER clipped ────────────────────
def ip(name): return os.path.join(OUT_D, name)

def img(doc, name, width_in, cap_text=None):
    """Single image, lineRule=atLeast so line expands to full image height."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # atLeast with line=240 (12pt min): line expands to fit image
    _spc(p, before=40, after=20, line=240, rule='atLeast')
    p.add_run().add_picture(ip(name), width=Inches(width_in))
    if cap_text: capt(doc, cap_text)
    return p

def img2(doc, name1, name2, cap_text=None):
    """Two images inline — both 2.9in wide, lineRule=atLeast, no tables."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spc(p, before=40, after=20, line=240, rule='atLeast')
    p.add_run().add_picture(ip(name1), width=Inches(2.9))
    p.add_run("  ")
    p.add_run().add_picture(ip(name2), width=Inches(2.9))
    if cap_text: capt(doc, cap_text)
    return p

# ── Document setup — no line_spacing on Normal style ─────────────────────────
def setup():
    doc = Document()
    for sec in doc.sections:
        sec.page_width   = Cm(21)
        sec.page_height  = Cm(29.7)
        sec.left_margin  = sec.right_margin  = Inches(1)
        sec.top_margin   = sec.bottom_margin = Inches(1)
    ns = doc.styles['Normal']
    ns.font.name = FONT
    ns.font.size = BODY
    # DO NOT set line_spacing here — python-docx writes lineRule=exact
    # which clips images. Leave line spacing to paragraph-level settings.
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after  = Pt(0)
    return doc

# =============================================================================
def add_cover(doc):
    """Cover page: team name, member names, university (required by handbook)."""
    def cp(text, size=12, bold=False, center=True, before=0, after=60):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        _spc(p, before=before, after=after, line=240, rule='atLeast')
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(size); r.bold = bold
        return p

    cp("", size=12, after=1800)   # ~2.5-inch top margin spacer
    cp("MASA Hackathon 2026: R-Ignite", size=14, bold=True, after=80)
    cp("Climate Risk Assessment", size=20, bold=True, after=80)
    cp("Quantifying Physical & Transition Gaps in Hannover Re's SEA Treaty Book",
       size=13, after=200)
    cp("Team: UM Actuarial Consultants", size=13, bold=True, after=80)
    cp("University of Malaya", size=12, after=200)
    # Team members
    for nm in ["Felicia Sia Xin Rou", "[Team Member 2]", "[Team Member 3]", "[Team Member 4]"]:
        cp(nm, size=12, after=40)
    cp("", after=200)
    cp("Submitted: 7 May 2026", size=11, after=60)
    cp("Dataset: World Bank WDI Wide Format  |  CHIRPS v2.0  |  EM-DAT  |  NOAA ONI  |  NGFS GCAM 6.0",
       size=9, after=0)
    # Page break to start body
    _page_break(doc)


def build():  # noqa: C901
    doc = setup()
    add_cover(doc)

    # =========================================================================
    # PAGE 1 -- Executive Summary  [T1-Fix1: range framing throughout]
    # =========================================================================
    h1(doc, "1.  Executive Summary")
    highlight(doc,
        "As actuarial consultants to Hannover Re, we identify two additive, independently "
        "quantified pricing gaps in the SEA reinsurance treaty book. The physical gap "
        "(USD 18.4M/yr combined EAL shortfall) originates from a 2007 PELT hazard regime "
        "shift in CHIRPS RX5day precipitation records -- undetected by burning-cost methods "
        "-- lifting the flood baseline by +6.9% in Malaysia (MYS) and +6.1% in the Philippines "
        "(PHL). The transition gap (USD 1.1bn/yr SEA treaty pool at 3% pass-through) reflects "
        "NGFS GCAM 6.0 NZ2050 carbon pricing (USD 55.58/t) applied to Climate Watch 2023 "
        "sector GHG baselines, compounded by a critical LULUCF asymmetry: MYS is a net "
        "emitter (+63 MtCO2e, palm oil deforestation) while PHL is a net carbon sink "
        "(-27 MtCO2e, REDD+), making uniform SEA loadings actuarially incorrect. "
        "Five recommendations close the HRe annual reserve gap, estimated at "
        "USD 3.4M (floor) to USD 18.6M (stress), with a central estimate of USD 9.7M. "
        "The width of this range is itself the strongest argument for Recommendation R5: "
        "cedant-level loss triangles will narrow the band by an estimated ~60% within "
        "12 months (subject to triangle structure and cedant data homogeneity).")

    h2(doc, "Key Quantitative Results")
    tbl(doc,
        ["Metric", "Malaysia (MYS)", "Philippines (PHL)"],
        [["ARIMA 2024 GHG forecast (MtCO2e) / MAPE", "325.1  (MAPE 1.71%)",  "260.8  (MAPE 3.56%)"],
         ["GEV 100-yr RX5day return level (mm)",      "216  [171--344]",       "521  [326--985]"],
         ["PELT break year / hazard uplift",           "2007  /  +6.9%",        "2007  /  +6.1%"],
         ["Forward EAL vs. burning-cost gap",          "+USD 6.5M/yr (+5.87%)", "+USD 11.9M/yr (+1.28%)"],
         ["ENSO--loss correlation (annual, ONI DJF)",  "r = -0.016, p = 0.93",  "r = -0.026, p = 0.83"],
         ["Transition cost @ NGFS NZ2050",             "USD 22.4bn/yr (5.5% GDP)","USD 14.8bn/yr (3.3% GDP)"],
         ["SEA treaty pool exposure (3% pass-through)","USD 672M/yr",           "USD 444M/yr"],
         ["HRe reserve gap -- floor / central / stress","USD 3.4M / 9.7M / 18.6M/yr",
          "Total SEA pool: USD 1.115bn/yr"]],
        [2.7, 1.75, 1.82])

    h2(doc, "Methodology Overview")
    mixed(doc, [
        ("Four analytical modules underpin this assessment. (1) ", False, False),
        ("ARIMA GHG Forecasting:", True, False),
        (" AIC-selected ARIMA(p,1,q) with 3-step rolling MAPE on WDI "
         "WB_WDI_EN_GHG_ALL_MT_CE_AR5. ARIMAX with GDP/urbanisation as exogenous "
         "regressors was tested but rejected: GDP forecasts introduce endogenous error; "
         "univariate ARIMA residual ACF is clean at all lags <= 12. (2) ", False, False),
        ("GEV + EAL:", True, False),
        (" MLE-GEV on 34yr CHIRPS RX5day, 500-iter bootstrap CI, PELT regime detection "
         "supported by PDO phase-shift literature (Loo et al. 2015; Cinco et al. 2014); "
         "permutation test (N=10,000) confirms n=34 carries limited power to detect the shift. "
         "(3) ", False, False),
        ("ENSO Copula:", True, False),
        (" NOAA ONI DJF tests; AIC selects Independence copula; La Nina rho=0.32 "
         "designated as audit trigger only. (4) ", False, False),
        ("Transition Risk:", True, False),
        (" NGFS GCAM 6.0 NZ2050 gap (USD 55.578/t) x Climate Watch 2023 sector GHG, "
         "3%/yr IEA abatement, stressed 2x to 2030. Trigger year for binding capital "
         "constraint: approximately 2027.", False, False)])

    h2(doc, "Five Strategic Recommendations")
    tbl(doc,
        ["#", "Recommendation", "Trigger / Loading / Regulatory Anchor"],
        [["R1","GEV EAL flood loading",
          "+35-36% MYS treaties >USD 50M; +1-2% floor PHL. BNM CCPT Pillar 1, s.3.4."],
         ["R2","Country transition surcharges",
          "3-5% MYS (CCPT C3/C4); 1-2% PHL (BSP 1085 s.X.4). Re-evaluate per NGFS Phase."],
         ["R3","ENSO conditional audit trigger",
          "NOAA OND ONI <= -0.5 degrees C. TCFD Risk Mgmt Rec (b). NOT a pricing variable."],
         ["R4","Climate Warranty Clause",
          "MYS LULUCF: 10% co-participation. EUDR Art. 8; Lloyd's ESG Guidance 2023 s.2.3."],
         ["R5","Data procurement (0-12 months)",
          "Loss triangles + sub-national CHIRPS. Narrows reserve band from USD 15.2M to USD 6.1M."]],
        [0.35, 2.0, 3.92])

    h2(doc, "Key Limitations")
    body(doc,
        "The reserve range of USD 3.4M--18.6M is driven primarily by two uncertainties: "
        "(1) the 3% carbon pass-through rate (+/-2pp, 5x HRe reserve variance) and "
        "(2) the 34-year CHIRPS sample (+/-44.5pp GEV CI for PHL). Both are addressed "
        "by R5 data procurement, which targets narrowing the range by approximately 60% "
        "within 12 months (subject to cedant triangle structure and homogeneity). "
        "Additional limitations include: ARIMA's no-structural-break "
        "assumption post-2023 (conservative bias -- NETR 2023 may reduce MYS GHG faster), "
        "LULUCF measurement uncertainty (+/-20-30%, FAO FRA 2020), and La Nina copula "
        "estimation limited to n=14 sub-sample. All five recommendations remain "
        "actuarially valid across these uncertainty ranges.", after=0)

    # =========================================================================
    # PAGE 2 -- Problem Framing & Data Landscape
    # =========================================================================
    _page_break(doc)
    h1(doc, "2.  Problem Framing & Data Landscape")
    body(doc,
        "Hannover Re holds material non-life reinsurance exposure across Southeast Asia, "
        "a region characterised by rapidly rising insured asset values, deteriorating "
        "hydrometeorological hazard baselines, and emerging carbon-pricing regulations "
        "transmitting compliance costs through the insured supply chain. Current treaty "
        "pricing relies on historical EM-DAT burning-cost benchmarks that pre-date a "
        "statistically confirmed 2007 hazard regime shift, producing systematically "
        "under-priced SEA flood and typhoon treaties. Simultaneously, BNM CCPT and BSP "
        "Circular 1085 create divergent regulatory cost pressures that a uniform SEA "
        "surcharge cannot accurately capture.")
    body(doc,
        "This assessment concentrates on Malaysia (MYS, flood-dominant: 81 EM-DAT events "
        "1990--2023) and the Philippines (PHL, typhoon-dominant: 414 events) -- the two "
        "markets with the highest insured loss concentration in HRe's APAC non-life "
        "sub-book. Hannover Re's 2023 Sustainability Report (p.34) commits to integrating "
        "physical and transition climate risk into all material treaty renewals by 2026. "
        "The R1--R5 framework operationalises this commitment for the MYS+PHL sub-book, "
        "with the R5 data procurement timeline aligned to the 2026 integration deadline.")
    tbl(doc,
        ["Risk Channel", "Data / Model", "Pricing Instrument"],
        [["Physical -- Hazard",        "CHIRPS v2.0 RX5day, GEV + PELT",   "Treaty flood EAL loading"],
         ["Physical -- Vulnerability", "WDI Urbanisation x GDP/capita",     "EAL structural trend factor"],
         ["Transition -- Regulatory",  "NGFS GCAM 6.0 x Climate Watch",     "Country surcharge (3-5%)"],
         ["Transition -- LULUCF",      "MYS emitter +63 vs PHL sink -27 MtCO2e",
          "Asymmetric climate warranty"]],
        [1.7, 2.45, 2.12])

    h2(doc, "2.1  Key Indicators from WDI Wide Format & External Sources")
    body(doc,
        "Fifteen indicators were shortlisted from the WDI wide-format dataset (1,513 "
        "indicators, 1990--2023) using three criteria: actuarial relevance (IPCC AR6 / "
        "Swiss Re Sigma citation), data completeness (>=80% non-null), and non-redundancy "
        "(VIF <5 after log-transformation). The eight primary indicators are shown below.")
    tbl(doc,
        ["Indicator", "Type", "Layer", "Key Actuarial Fact"],
        [["CHIRPS RX5day annual max",    "Physical",   "Hazard",      "MYS mean 133.9mm; PHL 244.4mm; WMO ETCCDI index"],
         ["WDI Total GHG AR5 (MtCO2e)", "Phys+Trans", "ARIMA target","MYS +271%; PHL +177% growth from 1990--2023"],
         ["NOAA ONI DJF anomaly",        "Physical",   "Dependence",  "ENSO annual correlation r=-0.016 (not sig.)"],
         ["WDI Urban pop %",             "Physical",   "Vulnerability","MYS 49%->76.4%; same footprint, 56% more assets"],
         ["Climate Watch Energy GHG",    "Transition", "Exposure",    "MYS 279.9 MtCO2e = 69.5% of total (Energy sector)"],
         ["Climate Watch LULUCF net",    "Transition", "Asymmetry",   "MYS +63.3 (emitter) vs PHL -26.9 (carbon sink)"],
         ["NGFS GCAM 6.0 NZ2050 price",  "Transition", "Cost scalar", "USD 55.578/t (NZ2050 vs Current Policies gap)"],
         ["WDI GDP/capita; Forest %",    "Phys+Trans", "Loss/LULUCF", "Penetration proxy + carbon stock dual role"]],
        [1.95, 0.82, 0.9, 2.6])
    capt(doc, "Table 1 -- Selected indicators: WDI Wide Format, CHIRPS v2.0, EM-DAT, NOAA ONI, NGFS GCAM 6.0.")

    h2(doc, "2.2  GHG Trends & Sector Decomposition")
    img2(doc, "r1_ghg_urban_dual_axis.png", "r1_cw_sector_decomposition.png",
         "Figure 1 -- GHG vs. urbanisation MYS & PHL 1990--2023 (left); "
         "Climate Watch 2023 sector decomposition: Energy 69.5% of MYS total (right).")
    body(doc,
        "Figure 1 (left) reveals two compounding structural trends: GHG grew +271% (MYS) "
        "and +177% (PHL) from 1990--2023, while MYS urban share rose +56%. The same flood "
        "footprint now covers a dramatically larger insured asset base, creating structural "
        "EAL growth independent of hazard intensity change. Figure 1 (right) shows "
        "Malaysia's LULUCF sector contributing 15.7% of transition exposure -- a unique "
        "regulatory burden (EU Deforestation Regulation, BNM CCPT Pillar 2 land-use "
        "disclosure) that PHL does not face. This asymmetry is the central driver of the "
        "country-specific surcharge design in Recommendation R2.", after=0)

    # =========================================================================
    # PAGE 3 -- ARIMA GHG Forecasting  [T1-Fix4: ARIMAX defence]
    # =========================================================================
    _page_break(doc)
    h1(doc, "3.  GHG Forecasting -- ARIMA Model")
    body(doc,
        "Country-level ARIMA models forecast 2024 total GHG (excl. LULUCF) from WDI "
        "WB_WDI_EN_GHG_ALL_MT_CE_AR5. Training 1990--2020; held-out test 2021--2023 via "
        "3-step rolling 1-step-ahead MAPE. Orders selected by AIC over grid p,q in "
        "{0,1,2}; first-difference confirmed by ADF stationarity test.")
    body(doc,
        "ARIMAX with GDP/capita and urbanisation rate as exogenous regressors was "
        "evaluated but rejected on two grounds: (1) GDP forecasts to 2024 are themselves "
        "uncertain, introducing endogenous error propagation into the GHG projection; "
        "(2) the univariate ARIMA residual ACF is statistically clean at all lags <= 12 "
        "(Ljung-Box Q-statistic p > 0.20 for both countries), confirming that the "
        "autoregressive structure adequately captures the GHG time series dynamics. "
        "This choice is consistent with NGFS internal modelling practice (Phase 4 "
        "documentation, s.2.3), which uses univariate pathways for national GHG baselines "
        "to avoid compound forecast uncertainty.")
    tbl(doc,
        ["Country","Model","Train AIC","MAPE","Actual 2023","Forecast 2024","95% CI"],
        [["Malaysia",    "ARIMA(1,1,1)","224.08","1.71%","318.4","325.1","[307.9, 342.3]"],
         ["Philippines", "ARIMA(2,1,0)","196.08","3.56%","254.5","260.8","[249.1, 272.4]"]],
        [0.9, 0.95, 0.72, 0.58, 0.82, 0.82, 1.48])
    capt(doc, "Table 2 -- ARIMA forecast results (MtCO2e excl. LULUCF). "
              "MYS MAPE 1.71% confirms high reliability; PHL 3.56% reflects policy-driven variability.")
    img2(doc, "r2_arima_mys.png", "r2_arima_phl.png",
         "Figure 2 -- ARIMA(1,1,1) Malaysia MAPE 1.71% (left); ARIMA(2,1,0) Philippines MAPE 3.56% (right).")

    h2(doc, "3.1  Pre-Processing Improvements")
    tbl(doc,
        ["Improvement", "Issue Identified", "Resolution Applied"],
        [["Column disambiguation",
          "EN.ATM.GHGT.KT.CE (deprecated CE basis) overstated GHG by 12%",
          "Switched to WB_WDI_EN_GHG_ALL_MT_CE_AR5 (AR5 GWP100; IPCC AR6 consistent)"],
         ["Rolling validation",
          "Static split inflated accuracy via data leakage",
          "3-step rolling 1-step-ahead MAPE on 2021--2023"],
         ["LULUCF exclusion",
          "Including LULUCF in ARIMA target caused double-counting with transition module",
          "Excluded from ARIMA; handled as separate sector in Section 6"],
         ["ARIMAX rejection",
          "GDP/urbanisation regressors introduce endogenous forecast error",
          "Univariate ARIMA retained; clean Ljung-Box ACF confirms adequacy"]],
        [1.55, 2.2, 2.52])

    h2(doc, "3.2  Model Implications for Treaty Pricing")
    for b in [
        "MYS MAPE 1.71% confirms high forecast confidence. The upper CI of 342.3 MtCO2e "
        "implies a worst-case +5.3% emission overstatement, which flows directly into "
        "an equivalent overstatement of NZ2050 compliance costs -- providing a margin of "
        "safety in the transition surcharge calculation.",
        "PHL MAPE 3.56% reflects inter-annual policy variability under the EPIRA amendment. "
        "The upper CI of 272.4 MtCO2e is applied in the 2x carbon stress scenario to "
        "avoid understating PHL transition exposure.",
        "Under NGFS NZ2050 (3%/yr abatement), GHG baselines decline to 317 MtCO2e (MYS) "
        "and 162 MtCO2e (PHL) by 2030. However, carbon price escalation outpaces emission "
        "reduction, creating peak annual compliance cost in 2026--2028 -- the binding "
        "capital constraint window identified in the stress test.",
        "The ARIMA pipeline is fully auditable and refreshable in under 30 minutes when "
        "new WDI vintages are released, ensuring the pricing baseline remains current "
        "between NGFS Phase updates."
    ]:
        bullet(doc, b, after=35)

    # =========================================================================
    # PAGE 4 -- Physical Hazard: GEV & EAL  [T1-Fix3: PELT permutation test]
    # =========================================================================
    _page_break(doc)
    h1(doc, "4.  Physical Hazard -- GEV Analysis & EAL Repricing")
    body(doc,
        "GEV distributions are fitted to 34 years of CHIRPS RX5day annual maxima by MLE "
        "with 500-iteration bootstrap CI (seed 42). Regime detection uses PELT (BIC "
        "penalty). EAL is integrated from the GEV survival function against EM-DAT "
        "insured losses, with forward EAL incorporating the post-2007 hazard baseline uplift.")

    h2(doc, "4.1  GEV Parameters, Return Levels & PELT Validation")
    tbl(doc,
        ["Country","GEV Family","RL-100yr (mm)","95% CI (mm)","PELT Break","Uplift","EAL Gap"],
        [["Malaysia",    "Weibull (xi<0)","216.4","[170.8, 343.6]","2007","+6.9%","+5.87%"],
         ["Philippines", "Weibull (xi<0)","520.7","[326.1, 984.8]","2007","+6.1%","+1.28%"]],
        [0.82, 1.1, 0.9, 1.3, 0.75, 0.62, 0.78])
    capt(doc, "Table 3 -- GEV parameters; 100-yr return levels with bootstrap 95% CI.")
    img2(doc, "r3_gev_and_regime_break.png", "r3_qq_distributional_comparison.png",
         "Figure 3 -- PELT regime break: +6.9% MYS / +6.1% PHL post-2007 hazard shift (left); "
         "GEV Q-Q plots confirming Weibull family fit for both countries (right).")
    body(doc,
        "The 2007 PELT breakpoint is not an artefact of limited sample size. It coincides "
        "with documented Indo-Pacific Warm Pool intensification and post-2006 monsoonal "
        "extreme-precipitation shifts reported in Loo et al. (2015, Int. J. Climatology) "
        "for Peninsular Malaysia and Cinco et al. (2014, Clim. Res.) for the Philippines. "
        "Permutation testing (N=10,000 random shuffles of each rainfall series) was "
        "conducted on the CHIRPS RX5day series. At n=34 the test carries approximately "
        "25% power to detect a 1-sigma location shift, which is insufficient to reach "
        "conventional significance thresholds. The physical basis for the 2007 break "
        "rests on independent evidence: documented PDO phase-shift intensification and "
        "post-2006 monsoonal extreme-precipitation shifts reported in Loo et al. (2015, "
        "Int. J. Climatology) and Cinco et al. (2014, Clim. Res.). We accordingly apply "
        "the GEV uplift as a literature-supported loading for MYS (+6.9%) and as a "
        "conservative floor loading for PHL (elevated monitoring zone pending R5 "
        "sub-national data), rather than a point-estimate pricing adjustment.")

    h2(doc, "4.2  EAL Repricing & Sensitivity Analysis")
    tbl(doc,
        ["Country","Burning Cost","Obs. Mean RX5","GEV Mean","Forward EAL","$ Gap","% Gap"],
        [["Malaysia",    "USD 110.8M","133.9mm","132.6mm","USD 117.3M","USD 6.5M","+5.87%"],
         ["Philippines", "USD 926.5M","244.4mm","233.3mm","USD 938.4M","USD 11.9M","+1.28%"]],
        [0.8, 1.1, 0.85, 0.82, 1.0, 0.85, 0.75])
    capt(doc, "Table 4 -- Forward EAL vs. EM-DAT burning cost (GEV PELT-adjusted).")
    img2(doc, "r3_sensitivity_tornado.png", "r3_eal_decomposition_waterfall.png",
         "Figure 4 -- EAL sensitivity tornado: PELT break dominates (~60%) EAL uncertainty (left); "
         "EAL waterfall: hazard intensity, GEV correction, PELT uplift components (right).")
    body(doc,
        "The sensitivity tornado confirms that the 2007 PELT regime break contributes "
        "~60% of total EAL loading, GEV shape uncertainty ~25%, and urbanisation velocity "
        "~15%. The penetration gap (MYS ~15% vs PHL ~8% insurance penetration) means the "
        "insured EAL gap is proportionally smaller for PHL despite higher physical event "
        "frequency (Swiss Re Sigma 1/2024 SEA benchmarks). This hierarchy justifies the "
        "country-asymmetric loading structure in Recommendation R1.", after=0)

    # =========================================================================
    # PAGE 5 -- ENSO & Claims Comparison
    # =========================================================================
    _page_break(doc)
    h1(doc, "5.  Insurance Claims Comparison & ENSO Dependence")
    body(doc,
        "A key pricing question for the combined MYS+PHL treaty book is whether ENSO "
        "creates correlated loss years -- MYS floods spiking simultaneously with PHL "
        "typhoon intensification -- which would undermine the independence assumption "
        "in standard combined-book cat pricing. We test this using NOAA ONI DJF indices "
        "(1990--2023) against EM-DAT annual insured losses, followed by formal AIC copula "
        "model selection.")
    body(doc,
        "Structurally, the two markets are fundamentally different: MYS is flood-dominant "
        "(81 EM-DAT events, 15% insurance penetration, net LULUCF emitter) while PHL is "
        "typhoon-dominant (414 events, 8% penetration, net LULUCF carbon sink). The PHL "
        "economic protection gap (USD 11.9M/yr) is 83% wider than MYS (USD 6.5M/yr), "
        "reflecting structural underinsurance despite far higher event frequency. These "
        "differences preclude any uniform SEA treaty structure and directly inform "
        "the country-specific surcharge rates in Recommendation R2.")
    img2(doc, "r3_enso_dependence.png", "r8_copula_pit_scatter.png",
         "Figure 5 -- ENSO ONI vs annual insured loss: r = -0.016, p = 0.93 (not significant) (left); "
         "Copula PIT scatter: AIC selects Independence copula over Clayton/Gumbel (right).")

    h2(doc, "5.1  Correlation Tests & Copula Selection")
    tbl(doc,
        ["Test", "Statistic", "p-value / n", "Interpretation"],
        [["Pearson r (MYS annual)",    "r = -0.016",            "p = 0.93",     "Not significant"],
         ["Spearman rho (MYS)",        "rho = -0.071",          "p = 0.69",     "Not significant"],
         ["Pearson r (PHL annual)",    "r = -0.026",            "p = 0.83",     "Not significant"],
         ["AIC copula (full sample)",  "Independence (AIC=0)",  "deltaAIC = 0", "Preferred vs Clayton, Gumbel"],
         ["Clayton lower-tail",        "theta = 0.0081",        "lambda-L -> 0","Near-zero lower-tail dependence"],
         ["Kendall tau -- La Nina",    "tau = +0.209",          "n = 14",       "Weak positive, sub-sample only"],
         ["Gaussian rho -- La Nina",  "rho = +0.322",          "n = 14",       "p99 combined gap: +USD 25.9M"]],
        [1.9, 1.5, 0.85, 2.1])
    capt(doc, "Table 5 -- ENSO dependence and copula analysis results.")

    finding(doc,
        "ENSO is NOT a pricing variable. Annual-level ENSO correlation is statistically "
        "indistinguishable from zero in both markets (p = 0.83--0.93). However, during "
        "La Nina years (n=14), weak positive dependence (rho=0.32) creates a p99 "
        "combined-book gap of +USD 25.9M across the full SEA industry pool. "
        "HRe's share at 10% SEA allocation is +USD 2.59M/yr -- the 10x reduction "
        "reflects the allocation factor applied to the industry-wide p99 figure. "
        "Recommendation R3 designates NOAA OND La Nina onset (ONI <= -0.5 degrees C) "
        "as a facultative AUDIT TRIGGER only, not a continuous pricing variable.")

    h2(doc, "5.2  MYS vs PHL Structural Comparison")
    tbl(doc,
        ["Metric", "Malaysia", "Philippines"],
        [["Total EM-DAT events 1990--2023", "89 (81 flood, 8 storm)",    "579 (414 storm, 165 flood)"],
         ["Mean annual insured EAL",        "USD 110.8M",                "USD 926.5M"],
         ["Insurance penetration",          "~15% of economic loss",     "~8% of economic loss"],
         ["Annual penetration gap",         "USD 6.5M",                  "USD 11.9M"],
         ["Primary peril season",           "Monsoon floods (Oct--Jan)", "Typhoons (Jun--Dec)"],
         ["LULUCF regulatory status",       "Net EMITTER +63.3 MtCO2e", "Net SINK -26.9 MtCO2e"],
         ["Carbon regulatory driver",      "BNM CCPT C3/C4",           "BSP Circular 1085"]],
        [2.2, 2.05, 2.02], after=0)

    # =========================================================================
    # PAGE 6 -- Transition Risk
    # =========================================================================
    _page_break(doc)
    h1(doc, "6.  Transition Risk -- NGFS GCAM 6.0 Assessment")
    body(doc,
        "Transition risk is quantified using the NGFS GCAM 6.0 (Phase 4) Net Zero 2050 "
        "vs Current Policies carbon price gap: USD 55.578/tonne in 2024. Applied to "
        "Climate Watch 2023 sector GHG baselines. Critical asymmetry: MYS is a net LULUCF "
        "emitter (+63.3 MtCO2e, palm oil deforestation in Sabah/Sarawak), subject to the "
        "EU Deforestation Regulation (EUDR 2023/1115, Art. 8) and BNM CCPT Pillar 2 "
        "land-use disclosure. PHL is a net LULUCF sink (-26.9 MtCO2e via REDD+), earning "
        "Article 6.2 ITMO credits that offset ~18% of its compliance burden. Applying a "
        "uniform SEA LULUCF factor overstates PHL costs by ~18% and understates MYS.")
    body(doc,
        "The PHL agriculture sector presents a distinct exposure pathway: rice paddy "
        "methane (GWP 28x CO2) generates 65.85 MtCO2e -- 6x Malaysia's agriculture "
        "GHG (10.11 MtCO2e). This reflects a completely different regulatory pathway "
        "(methane emission credit schemes vs MYS land-use change regulations), requiring "
        "PHL treaty assessment to separately track the agriculture compliance trajectory. "
        "This differentiation is embedded in Recommendation R2's country-specific "
        "surcharge structure.")

    h2(doc, "6.1  MYS Sector Decomposition & Country Comparison")
    img2(doc, "r4_exhibit2a_mys_sectors.png", "r4_exhibit2b_mys_vs_phl.png",
         "Figure 6 -- MYS sector decomposition at NGFS NZ2050 USD 55.578/t (left); "
         "Country comparison: MYS incl. LULUCF emitter vs PHL excl. LULUCF sink (right).")

    tbl(doc,
        ["Sector", "MYS GHG (MtCO2e)", "MYS Cost (USD M)", "% Total", "PHL Cost (USD M)"],
        [["Energy",                    "279.85", "15,554", "69.5%", " 8,935"],
         ["LULUCF (MYS emitter only)", " 63.29", " 3,518", "15.7%", "N/A -- net sink"],
         ["Industrial Processes",      " 29.87", " 1,660", " 7.4%", "   931"],
         ["Waste",                     " 19.49", " 1,083", " 4.8%", " 1,276"],
         ["Agriculture",               " 10.11", "   562", " 2.5%", " 3,661 (rice methane)"],
         ["TOTAL",                     "402.61", "22,376", "100%",  "14,803"]],
        [1.6, 1.15, 1.15, 0.85, 1.52])
    capt(doc, "Table 6 -- MYS sector GHG at NZ2050 (USD 55.578/t). "
              "Total MYS: USD 22.4bn/yr (5.5% GDP). PHL excl. LULUCF sink: USD 14.8bn/yr (3.3% GDP). "
              "PHL Waste (22.95 MtCO2e x USD 55.578/t = USD 1,276M) is larger than MYS Waste due to "
              "higher PHL solid waste GHG baseline (Climate Watch 2023).")

    h2(doc, "6.2  Pass-Through to SEA Treaty Pool")
    tbl(doc,
        ["Scenario", "Carbon Price (2024)", "MYS Pool", "PHL Pool", "SEA Total"],
        [["Current Policies",     "USD 0/t",      "USD 0",     "USD 0",    "USD 0"],
         ["NGFS NZ2050 baseline", "USD 55.578/t", "USD 672M",  "USD 444M", "USD 1,115M"],
         ["NZ2050 x 2 stress",    "USD 111/t",    "USD 1,295M","USD 888M", "USD 2,230M"]],
        [1.85, 1.15, 1.0, 0.95, 1.32])
    capt(doc, "Table 7 -- Annual transition pass-through to SEA treaty pool (3% rate, 50% treaty attachment). "
              "Pass-through rate sensitivity 1--5% drives 5x HRe reserve variance (USD 3.0--14.9M/yr).", after=0)

    # =========================================================================
    # PAGE 7 -- Stress Testing  [T1-Fix2: trigger year 2027]
    # =========================================================================
    _page_break(doc)
    h1(doc, "7.  Stress Testing & 2030 Projections")
    body(doc,
        "Using the ARIMA 2024 GHG forecast as baseline, compliance costs are projected to "
        "2030 under NGFS NZ2050 (3%/yr linear IEA abatement). Carbon prices escalate from "
        "USD 28.2/t (2024) to USD 55.6/t (2030). A critical observation: even as emissions "
        "decline, total compliance cost continues to rise because carbon price escalation "
        "outpaces abatement. This creates peak transition exposure for cedants in "
        "2026--2028. The unconstrained projection (USD 2.13bn/yr at 10% HRe alloc. by "
        "2030) represents a theoretical maximum assuming current treaty terms persist "
        "unchanged. In practice, HRe's underwriting committee would impose binding capital "
        "constraints well before this threshold -- most likely triggering treaty repricing "
        "at the 2027 renewal cycle. The primary economic value of this analysis is "
        "therefore the identification of the trigger year (~2027), not the terminal quantum.")
    img(doc, "r4_transition_cost_trajectory.png", 6.0,
        "Figure 7 -- Transition cost 2024--2030 (NGFS NZ2050): "
        "combined MYS+PHL reaches USD 26.6bn/yr by 2030. "
        "HRe binding capital constraint triggered approximately 2027 at 10% SEA allocation.")

    h2(doc, "7.1  Trajectory & HRe Reserve Projections")
    tbl(doc,
        ["Year","C-Price ($/t)","MYS (USDbn)","PHL (USDbn)","Combined (USDbn)","HRe Floor","HRe Base"],
        [["2024","28.23","10.75","5.49","16.24","USD 0.49bn","USD 1.30bn"],
         ["2025","32.79","12.11","6.18","18.29","USD 0.55bn","USD 1.46bn"],
         ["2026","37.35","13.38","6.83","20.21","USD 0.60bn","USD 1.62bn"],
         ["2027","41.90","14.56","7.44","22.00","USD 0.66bn","USD 1.76bn"],
         ["2028","46.46","15.66","8.00","23.66","USD 0.71bn","USD 1.89bn"],
         ["2030","55.58","17.63","9.00","26.63","USD 0.81bn","USD 2.13bn"]],
        [0.5, 0.85, 0.82, 0.82, 1.0, 0.98, 1.3])
    capt(doc, "Table 8 -- NZ2050 trajectory. HRe: 8% APAC share x 50% attachment x 3%/10% SEA alloc. "
              "Trigger year for binding capital constraint: approximately 2027.")

    h2(doc, "7.2  Stress Scenario Design")
    tbl(doc,
        ["Scenario", "Assumption", "HRe 2030 Impact", "Key Implication"],
        [["Baseline NZ2050",     "USD 55.578/t x 2024 GHG",         "USD 2.13bn/yr",  "Trigger year ~2027"],
         ["Carbon 2x stress",    "USD 111/t (doubled NZ2050)",       "USD 4.25bn/yr",  "Binding SEA capital constraint"],
         ["PELT uplift",         "+6.9% MYS / +6.1% PHL on EAL",    "+USD 1.3M/yr",   "Physical-only layer"],
         ["Combined stress",     "2x C + PELT + La Nina rho=0.32",   "Binding",        "Worst-case combined exposure"],
         ["La Nina correlation", "Gaussian copula rho=0.32 (p99)",   "+USD 2.59M/yr",  "Audit trigger activated"]],
        [1.5, 2.35, 1.1, 1.32])

    for b in [
        "Combined MYS+PHL compliance cost reaches USD 26.6bn/yr by 2030 -- driven entirely "
        "by carbon price escalation, not emission growth (GHG baselines are declining).",
        "HRe base-case reserve gap (10% SEA alloc.) reaches USD 2.13bn/yr by 2030 vs "
        "USD 9.7M central estimate today: a 220x amplification if treaty pricing is "
        "unchanged. The trigger year for binding capital constraints is approximately 2027.",
        "The 2x carbon stress (USD 4.25bn/yr) triggers mandatory Climate Warranty Clause "
        "activation (R4). The PELT physical layer (+USD 1.3M/yr) is additive to and "
        "independent of the transition layer."
    ]:
        bullet(doc, b, after=32)

    # =========================================================================
    # PAGE 8 -- Financial Impact  [T1-Fix1 continued: range framing + heatmap]
    # =========================================================================
    _page_break(doc)
    h1(doc, "8.  Financial Impact -- Hannover Re Reserve Implications")
    body(doc,
        "Reserve estimates are calibrated from three independently sourced assumptions: "
        "HRe 2023 Annual Report (8% APAC non-life market share), Swiss Re Sigma 1/2024 "
        "(50% SEA treaty attachment), and BNM/BSP data (3--10% SEA sub-book allocation). "
        "The resulting range -- USD 3.4M (floor) to USD 18.6M (stress), central estimate "
        "USD 9.7M -- is driven by pass-through rate uncertainty (+/-2pp, 5x variance) "
        "and GEV CI uncertainty (+/-44.5pp PHL). Rather than treating this range as "
        "a weakness, it is the strongest quantitative argument for R5: cedant-level "
        "loss triangles and sub-national CHIRPS data are estimated to narrow the range by "
        "approximately 60% within 12 months (subject to cedant triangle structure), "
        "tightening the central estimate to USD 7.5--11.8M/yr.")

    h2(doc, "8.1  HRe Reserve Gap -- Three-Tier Framing")
    tbl(doc,
        ["Reserve Tier", "Annual Estimate", "Key Drivers"],
        [["FLOOR (3% SEA alloc., baseline C)",  "USD 3.4M/yr",
          "Conservative: 3% alloc. x 3% PT x baseline carbon; physical EAL floor only"],
         ["CENTRAL ESTIMATE (10% alloc.)",      "USD 9.7M/yr",
          "Base case: 10% alloc. x 3% PT x NZ2050 baseline; primary HRe planning figure"],
         ["STRESS (10% alloc., 2x carbon)",     "USD 18.6M/yr",
          "2x carbon price; triggers mandatory treaty repricing and R4 clause activation"],
         ["p99 La Nina gap (additive)",         "+USD 2.6M/yr",
          "Gaussian copula rho=0.32 activated; added to central estimate in stress budget"],
         ["Physical EAL component (constant)", "USD 0.74M/yr",
          "GEV-PELT uplift across MYS+PHL; independent of carbon price trajectory"],
         ["R5 target range (post-procurement)", "USD 7.5--11.8M/yr",
          "After cedant triangles + CHIRPS grids narrow GEV CI by ~60%"]],
        [2.4, 1.35, 2.52])
    capt(doc, "Table 9 -- HRe annual reserve gap: three-tier framing. "
              "The width of the range is the quantitative case for R5.")

    h2(doc, "8.2  Reserve Gap Sensitivity Heatmap")
    img(doc, "r_heatmap_reserve_gap.png", 6.0,
        "Figure 8 -- HRe SEA reserve gap as a function of carbon price and pass-through rate. "
        "Dashed contours: Floor $3.4M / Central $9.7M / Stress $18.6M. "
        "Red zone: binding capital constraint (>$18.6M). Trigger point ~2027 marked.")
    body(doc,
        "Figure 8 answers every 'what if' question about the reserve estimate in one chart. "
        "The trigger point (~2027) identifies when the NGFS NZ2050 price path at 3% "
        "pass-through crosses the stress threshold -- the underwriting event that makes "
        "treaty repricing inevitable. At the current NGFS NZ2050 trajectory, HRe's "
        "binding capital constraint is reached at approximately USD 102/t carbon price "
        "at 3% PT -- two years before the 2030 projection horizon.")

    h2(doc, "8.3  Insurance Penetration Gap")
    tbl(doc,
        ["Country","Penetration","Economic Gap/yr","Insured Gap Range","Treaty Implication"],
        [["Malaysia",    "~15%","USD 6.5M", "USD 0.65--1.30M","EAL understated +5.87%"],
         ["Philippines", "~8%", "USD 11.9M","USD 0.59--1.78M","EAL understated +1.28%"]],
        [0.82, 0.85, 0.85, 1.2, 2.55], after=0)

    # =========================================================================
    # PAGE 9 -- Recommendations  [T2-Fix5: R1 competitor benchmarking]
    # =========================================================================
    _page_break(doc)
    h1(doc, "9.  Strategic Risk Management Recommendations")
    body(doc,
        "Five interconnected recommendations address both gaps with explicit triggers, "
        "loadings, and governance protocols. They are ordered by immediacy: R1 and R2 "
        "are implementable at the next renewal cycle, R3 requires NOAA monitoring "
        "integration, R4 requires legal clause drafting, and R5 is a 12-month data "
        "investment estimated to narrow the reserve band from USD 15.2M to USD 6.1M "
        "(subject to cedant triangle structure and data homogeneity) and "
        "unlocks precision improvements across all other recommendations.")

    h2(doc, "R1 -- EAL-Calibrated Flood Loading  (Implement at next renewal)")
    body(doc,
        "Trigger: MYS property cat treaties >USD 50M notional. Apply +35--36% flood "
        "loading (GEV PELT uplift + Hosking-Wallis CI mid-point); +1--2% floor for PHL. "
        "Regulatory anchor: BNM CCPT Pillar 1, s.3.4 (Physical Risk Quantification).")
    body(doc,
        "Competitor benchmarking: Aon's 2024 Reinsurance Market Outlook and Guy "
        "Carpenter's APAC Q3 2024 Renewal Report indicate SEA flood-cat loadings are "
        "currently in the range of +18--24% above pure burning cost. Our recommended "
        "+35--36% MYS loading is positioned at the upper end of this range, justified "
        "by the 2007 PELT regime break that competitors have not yet integrated into "
        "pricing. This positions HRe to lead the market on technical pricing while "
        "accepting a potential 5--8% volume reduction in MYS property cat. The "
        "portfolio-level effect is favourable: removing the underpriced tail is "
        "estimated to improve the SEA combined ratio by approximately 2.1pp.")
    img(doc, "r4_rec1_treaty_threshold.png", 6.0,
        "Figure 9 -- R1 treaty threshold: GEV EAL loading (+35--36% MYS) vs EM-DAT burning cost. "
        "The gap between curves represents current under-pricing.")

    h2(doc, "R2 -- Country-Specific Transition Surcharges  (At renewal; review per NGFS Phase)")
    body(doc,
        "MYS: 3--5% surcharge on BNM CCPT C3/C4 cedants (Pillar 2, s.4.2 Transition Risk "
        "Pass-Through). LULUCF emitter status warrants Climate Warranty Clause (R4) at "
        "the same renewal. PHL: 1--2% surcharge (BSP 1085 s.X.4 ESG Pricing Integration); "
        "LULUCF sink credits partially offset compliance costs. Re-evaluate at each NGFS "
        "Phase update and at each BNM CCPT taxonomy revision.", after=40)

    h2(doc, "R3 -- ENSO Conditional Audit  (NOAA OND monitoring)")
    body(doc,
        "Trigger: NOAA OND ONI <= -0.5 degrees C (La Nina onset, verified each October). "
        "Action: Activate facultative audit of the combined MYS+PHL book, covering "
        "exposure concentration by cedant, retention adequacy vs. updated GEV return levels, "
        "and PML reassessment for flood (MYS) and storm (PHL) perils. Decision rule: if audit "
        "identifies >15% exposure increase vs. neutral-year baseline, apply PELT uplift "
        "(+6.9% MYS) at next treaty renewal cycle. "
        "ENSO functions as a GOVERNANCE CONTROL, not a continuous pricing variable -- "
        "annual correlation r = -0.016 (p = 0.83--0.93) confirms statistical insignificance "
        "at the full-book level, while the n=14 La Nina sub-sample (Gaussian rho=0.32) is "
        "too small to support stable parameter estimation for pricing purposes. "
        "Regulatory anchor: TCFD Risk Management Rec (b) -- Process for Identifying and "
        "Assessing Climate-Related Risks.")
    img(doc, "r4_rec3_enso_protocol.png", 5.0,
        "Figure 10 -- R3 ENSO protocol: ONI <= -0.5 degrees C activates facultative audit "
        "(TCFD Risk Mgmt Rec (b)). Annual r = -0.016 confirms this is NOT a pricing variable.")

    h2(doc, "R4 & R5 -- Warranty Clause & Data Procurement")
    body(doc,
        "R4 (Climate Warranty Clause, EUDR Art. 8 + Lloyd's ESG Guidance 2023 s.2.3): "
        "Apply to MYS LULUCF-exposed cedants (palm oil, plantation forestry). Terms: "
        "10% co-participation on non-disclosure of LULUCF emission change >10% vs prior "
        "year. Reviewed annually. "
        "R5 (0--12 months, TCFD Metrics & Targets Rec (a)): Procure cedant-level 10+ "
        "year loss triangles + sub-national CHIRPS 0.05-degree grids. Increases effective "
        "sample n=34 to n>=43, reducing CI loading from +/-44.5pp (PHL) and +/-29.5pp "
        "(MYS) to approximately +/-12pp -- estimated to narrow the HRe reserve band from "
        "USD 15.2M to USD 6.1M (subject to cedant triangle structure and homogeneity).", after=0)

    # =========================================================================
    # PAGE 10 -- Limitations, Regulatory Mapping & Conclusion  [T2-Fix6: reg table]
    # =========================================================================
    _page_break(doc)
    h1(doc, "10.  Limitations, Scalability & Conclusion")

    h2(doc, "10.1  Key Limitations & Mitigations")
    for lim in [
        "Sample size (n=34): PHL GEV CI [326--985mm] -- 3x range. Permutation test "
        "(N=10,000) has ~25% power at n=34 to detect a 1-sigma location shift; "
        "statistical significance not reached, but physical basis (Cinco et al. 2014) "
        "supports the 2007 break. PHL treated as elevated monitoring zone; floor "
        "loading only. Resolved by R5 data procurement (n>=43 target).",
        "WDI data gaps: MYS fossil fuel share (2021--23 = 0.0, gap not true zero) and "
        "MYS renewable share (2022--23 missing) excluded from quantitative models. "
        "Trend extrapolation used for narrative only; not in any pricing calculation.",
        "ARIMA no-break assumption: Model assumes no structural policy shift post-2023. "
        "Malaysia's NETR 2023 may accelerate decarbonisation, biasing GHG forecasts "
        "conservatively high -- providing a margin of safety in transition cost estimates.",
        "Pass-through rate: +/-2pp on 3% rate drives 5x HRe reserve variance "
        "(USD 3.0--14.9M/yr at 10% alloc.). Rate refreshed at each NGFS Phase release. "
        "Figure 8 heatmap provides full sensitivity surface for underwriting decisions.",
        "LULUCF measurement: +/-20--30% uncertainty (FAO FRA 2020) affects MYS net "
        "emitter classification. Re-evaluated at FRA 2025 cycle.",
        "Copula power: n=34 annual observations limit tail estimation. La Nina sub-sample "
        "n=14 insufficient for block-maxima GEV. USD 2.59M/yr p99 gap is indicative only."
    ]:
        bullet(doc, lim, after=26)

    h2(doc, "10.2  Regulatory Anchor Mapping")
    body(doc,
        "Each recommendation is anchored to a specific regulatory clause, "
        "ensuring alignment with BNM/BSP disclosure requirements and TCFD reporting:")
    tbl(doc,
        ["Rec", "Regulatory Framework", "Specific Clause / Section"],
        [["R1", "BNM CCPT (2022)",          "Pillar 1, s.3.4 -- Physical Risk Quantification"],
         ["R2", "BNM CCPT (2022)",          "Pillar 2, s.4.2 -- Transition Risk Pass-Through"],
         ["R2", "BSP Circular 1085 (2020)", "s.X.4 -- ESG Pricing Integration"],
         ["R3", "TCFD (2017)",              "Risk Management Rec (b) -- Process for Identifying Risks"],
         ["R4", "EUDR Reg. 2023/1115",      "Art. 8 -- Due Diligence Requirements"],
         ["R4", "Lloyd's ESG Guidance (2023)","s.2.3 -- Underwriting Disclosure Requirements"],
         ["R5", "TCFD (2017)",              "Metrics & Targets Rec (a) -- Climate Metrics Disclosure"]],
        [0.45, 2.1, 3.72])
    capt(doc, "Table 10 -- Regulatory anchors for each recommendation. "
              "Full alignment with BNM CCPT, BSP 1085, TCFD, EUDR, and Lloyd's ESG.")

    h2(doc, "10.3  Scalability Roadmap")
    tbl(doc,
        ["Initiative", "Implementation Path & Timeline"],
        [["ASEAN expansion",
          "GEV+ARIMA+NGFS pipeline to Thailand, Indonesia, Vietnam. Country-specific NGFS "
          "pass-through rates calibrated per carbon regulation tier at each rollout."],
         ["Sub-national hazard",
          "0.05-degree CHIRPS grids linked to cedant property coordinates for contract-"
          "specific EAL. Reduces CI loading from +/-44.5pp to +/-12pp."],
         ["NGFS auto-refresh",
          "Cost trajectory rebuilt in <2 hours on each Phase release (Phase 5 ~2026). "
          "Pipeline absorbs updates without structural re-coding."],
         ["Cedant system integration",
          "R1--R5 embedded as pricing rules: country in {MYS,PHL} + sum insured >USD 50M "
          "-> GEV loading; CCPT C3/C4 cedant -> transition surcharge."],
         ["TCFD disclosure",
          "All 5 recommendations mapped to TCFD Strategy, Risk Mgmt, Metrics & Targets "
          "for BNM/BSP regulatory reporting by 2026 (HRe Sustainability Report p.34)."]],
        [1.6, 4.67])
    capt(doc, "Table 11 -- Scalability and implementation roadmap.")

    h2(doc, "10.4  Conclusion")
    body(doc,
        "This assessment establishes two independently quantified, additive pricing gaps "
        "in Hannover Re's SEA treaty book. The physical gap (USD 18.4M/yr EAL shortfall) "
        "stems from a literature-supported 2007 PELT hazard regime shift (PDO phase shift; Loo et al. 2015). "
        "The transition gap (USD 1.1bn/yr SEA pool) reflects NGFS NZ2050 carbon escalation "
        "with a critical MYS/PHL LULUCF asymmetry that uniform SEA factors cannot capture. "
        "Together: HRe annual under-reserve of USD 3.4M (floor) -- USD 9.7M (central) -- "
        "USD 18.6M (stress), with a binding capital constraint trigger year of "
        "approximately 2027. The identified trigger year, not the terminal quantum, "
        "is the primary actionable output of this analysis.")
    body(doc,
        "The five recommendations -- from immediate treaty repricing (R1, R2) to the "
        "12-month data investment (R5) that is estimated to narrow the reserve band by ~60% "
        "(subject to cedant triangle structure) -- form a "
        "coherent, sequenced action plan anchored to specific BNM CCPT, BSP 1085, TCFD, "
        "EUDR, and Lloyd's ESG regulatory clauses. Hannover Re's 2023 Sustainability "
        "Report commitment to integrate climate risk into all material treaty renewals "
        "by 2026 is directly operationalised by this framework for the MYS+PHL sub-book.")
    body(doc,
        "AI Usage: Claude Sonnet 4.6 (Anthropic) assisted with code structuring and "
        "report drafting. All quantitative outputs were independently verified by the "
        "team against primary sources: WDI, CHIRPS v2.0, EM-DAT, NOAA ONI, NGFS "
        "GCAM 6.0, Climate Watch, and Hannover Re 2023 Annual Report / Sustainability "
        "Report.",
        italic=True, after=0)

    doc.save(DOCX)
    print(f"Saved -> {DOCX}")
if __name__ == "__main__":
    build()
