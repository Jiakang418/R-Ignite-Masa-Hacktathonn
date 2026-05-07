"""
Generate final documentation Word document for R-Ignite Hackathon 2026.
Times New Roman, 12pt, with all output images embedded.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/Users/a1357/Documents/GitHub/R-Ignite-Masa-Hacktathonn/outputs"
SAVE_PATH = "/Users/a1357/Documents/GitHub/R-Ignite-Masa-Hacktathonn/outputs/Documentation_Numbers_Final.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=12, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # force complex-script font too
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.insert(0, rFonts)


def add_paragraph(doc, text="", bold=False, italic=False, size=12,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size, color=color)
    return p


def add_heading(doc, text, level=1, size=None, color=None, space_before=12, space_after=6):
    sizes = {1: 14, 2: 12, 3: 12}
    sz = size or sizes.get(level, 12)
    bolds = {1: True, 2: True, 3: True}
    p = add_paragraph(doc, text, bold=bolds.get(level, True), size=sz,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      color=color or (0x1F, 0x49, 0x7D) if level == 1 else color,
                      space_before=space_before, space_after=space_after)
    return p


def add_image(doc, filename, caption, width=Inches(6.0)):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        add_paragraph(doc, f"[Image not found: {filename}]", italic=True, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = add_paragraph(doc, caption, italic=True, size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def add_table_row(table, cells_data, header=False, shading=None):
    row = table.add_row()
    for i, (cell_text, cell_obj) in enumerate(zip(cells_data, row.cells)):
        p = cell_obj.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(cell_text))
        run.bold = header
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rPr.insert(0, rFonts)
        if shading:
            tc = cell_obj._tc
            tcPr = tc.get_or_add_tcPr()
            sh = OxmlElement("w:shd")
            sh.set(qn("w:val"), "clear")
            sh.set(qn("w:color"), "auto")
            sh.set(qn("w:fill"), shading)
            tcPr.append(sh)
    return row


def make_table(doc, headers, rows, col_widths=None, header_shade="1F497D", header_font_color=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hrow = table.rows[0]
    for i, (h, cell) in enumerate(zip(headers, hrow.cells)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:color"), "auto")
        sh.set(qn("w:fill"), header_shade)
        tcPr.append(sh)
        if col_widths:
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(col_widths[i] * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
    for r_idx, row_data in enumerate(rows):
        new_row = table.add_row()
        shade = "D9E2F3" if r_idx % 2 == 0 else "FFFFFF"
        for i, (cell_text, cell) in enumerate(zip(row_data, new_row.cells)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            sh = OxmlElement("w:shd")
            sh.set(qn("w:val"), "clear")
            sh.set(qn("w:color"), "auto")
            sh.set(qn("w:fill"), shade)
            tcPr.append(sh)
    return table


def add_bullet(doc, text, level=0, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def page_break(doc):
    doc.add_page_break()


# ── document ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    add_paragraph(doc, "", space_after=24)
    add_paragraph(doc, "Quantifying Climate Risk & Pricing Adequacy in SEA",
                  bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=(0x1F, 0x49, 0x7D), space_after=6)
    add_paragraph(doc, "Documentation", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    add_paragraph(doc, "Team:", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, "Numbers", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # team table
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hrow = t.rows[0]
    for txt, cell in zip(["Name", "Email"], hrow.cells):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto")
        sh.set(qn("w:fill"), "1F497D")
        tcPr.append(sh)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    members = [
        ("Khe Jia Kang", "jiakangkhe@gmail.com"),
        ("Lean Wen Jie", "jetlean0707@gmail.com"),
        ("Lee Jing Xuan", "jingx349@gmail.com"),
        ("Lau Hiap Meng", "hiapmenglau@gmail.com"),
        ("Felicia Sia Xin Rou", "flc.066734@gmail.com"),
    ]
    for name, email in members:
        row = t.add_row()
        for txt, cell in zip([name, email], row.cells):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    add_paragraph(doc, "", space_after=6)
    add_paragraph(doc, "MASA Hackathon 2026 | Hannover Re | Universiti Malaya",
                  italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    page_break(doc)

    # ── DECLARATION ───────────────────────────────────────────────────────────
    add_heading(doc, "Declaration of Originality & Compliance", level=1)
    add_paragraph(doc, "To: The Organizing Committee, MASA Hackathon 2026", space_after=6)

    add_heading(doc, "1. Statement of Originality", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc,
        'We, the undersigned members of "Numbers", hereby declare that the work submitted '
        'for the MASA Hackathon 2026 is our original creation. We certify that:')
    add_bullet(doc,
        "The conceptual framework, methodology, and code were developed solely by the team members during the hackathon period.")
    add_bullet(doc,
        "The integration of Pruned Exact Linear Time (PELT) regime break detection with Generalized Extreme Value (GEV) theory, "
        "and the subsequent application of an EIOPA-anchored pass-through rate matrix for treaty repricing, represents an original, "
        "actuarially sound synthesis.")
    add_bullet(doc,
        "To the best of our knowledge, this work does not infringe upon the intellectual property rights of any third party.")

    add_heading(doc, "2. Compliance with Official Rules", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc,
        "We confirm that we have read, understood and adhered to all the official rules and regulations governing the MASA "
        "Hackathon 2026. This submission complies with all eligibility requirements, data usage restrictions, and submission deadlines.")

    add_heading(doc, "3. Disclosure of External Resources & AI Tools", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc, "In the interest of transparency and in accordance with the competition guidelines:")
    add_bullet(doc, "All external datasets and academic references used in this project have been properly cited.")
    add_bullet(doc,
        "All AI tools used for assistance in code optimization, documentation, or brainstorming have been fully disclosed "
        "and documented in the AI_Usage.md and README.md files included in this submission.")
    add_bullet(doc,
        "The core actuarial logic and the novel synthesis of PELT and GEV theories remain the primary intellectual output of the team.")
    add_paragraph(doc,
        "We affirm that this submission was developed with the integrity and professional care expected within the actuarial profession. "
        "We accept full responsibility for the content and conclusions presented in our submission.")

    add_heading(doc, "Team Members & Signatures:", level=2, color=(0, 0, 0))
    sigs = [
        ("Khe Jia Kang", "Jia Kang"),
        ("Lean Wen Jie", "Wen Jie"),
        ("Lee Jing Xuan", "Sean"),
        ("Lau Hiap Meng", "Hiap Meng"),
        ("Felicia Sia Xin Rou", "Felicia"),
    ]
    st = doc.add_table(rows=1, cols=2)
    st.style = "Table Grid"
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow2 = st.rows[0]
    for txt, cell in zip(["Name", "Signature"], hrow2.cells):
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt); run.bold = True
        run.font.name = "Times New Roman"; run.font.size = Pt(12)
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
        sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), "1F497D")
        tcPr.append(sh); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for name, sig in sigs:
        row = st.add_row()
        for txt, cell in zip([name, sig], row.cells):
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt); run.italic = (txt == sig); run.bold = (txt == sig)
            run.font.name = "Times New Roman"; run.font.size = Pt(12)
    page_break(doc)

    # ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    toc = [
        ("1.", "Executive Summary", "4"),
        ("2.", "Problem Framing & Data Landscape", "5"),
        ("   2.1", "Key Indicators (WDI Wide Format + External Sources)", "5"),
        ("   2.2", "GHG Trends & Sector Decomposition", "5"),
        ("3.", "GHG Forecasting -- ARIMA Model", "6"),
        ("   3.1", "Pre-Processing Improvements", "6"),
        ("   3.2", "Model Implications for Treaty Pricing", "6"),
        ("4.", "Physical Hazard -- GEV Analysis & EAL Repricing", "7"),
        ("   4.1", "GEV Parameters & Return Levels & PELT Validation", "7"),
        ("   4.2", "EAL Repricing & Sensitivity Analysis", "7"),
        ("5.", "Insurance Claims Comparison & ENSO Dependence", "8"),
        ("   5.1", "Correlation Tests & Copula Selection", "8"),
        ("   5.2", "MYS vs PHL Structural Comparison", "9"),
        ("6.", "Transition Risk -- NGFS GCAM 6.0 Assessment", "9"),
        ("   6.1", "MYS Sector Decomposition & Country Comparison", "9"),
        ("   6.2", "Pass-Through to SEA Treaty Pool", "10"),
        ("7.", "Stress Testing & 2030 Projections", "10"),
        ("   7.1", "Trajectory & HRe Reserve Projections", "11"),
        ("   7.2", "Stress Scenario Design", "11"),
        ("8.", "Financial Impact -- Hannover Re Reserve Implications", "11"),
        ("   8.1", "HRe Reserve Gap -- Three-Tier Framing", "12"),
        ("   8.2", "Reserve Gap Sensitivity Heatmap", "12"),
        ("   8.3", "Insurance Penetration Gap", "12"),
        ("9.", "Strategic Risk Management Recommendations", "12"),
        ("10.", "Limitations, Scalability & Conclusion", "13"),
        ("   10.1", "Key Limitations & Mitigations", "13"),
        ("   10.2", "Regulatory Anchor Mapping", "14"),
        ("   10.3", "Conclusion", "14"),
    ]
    for num, title, pg in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}  {title}")
        run.font.name = "Times New Roman"; run.font.size = Pt(12)
        run.bold = not num.startswith("   ")
    page_break(doc)

    # ── SECTION 1: EXECUTIVE SUMMARY ─────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)
    add_paragraph(doc,
        "We identify two additive pricing gaps in HRe's SEA treaty book. (1) Physical gap (USD 18.4M/yr EAL shortfall) from a 2007 "
        "PELT regime shift undetected by burning-cost methods (+6.9% MYS / +6.1% PHL flood baseline). (2) Transition gap "
        "(USD 1.1bn/yr SEA pool at 3% pass-through) from NGFS NZ2050 carbon pricing (USD 55.58/t), with a critical LULUCF "
        "asymmetry: MYS net emitter (+63 MtCO₂e, palm oil) vs PHL net sink (-27 MtCO₂e, REDD+), making uniform SEA "
        "loadings actuarially incorrect. HRe annual reserve gap: USD 3.4M (floor) -- USD 9.7M (central) -- USD 18.6M (stress). "
        "The range width is the strongest argument for R5: cedant triangles narrow it ~60% in 12 months.")

    add_heading(doc, "Key Quantitative Results", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Metric", "Malaysia (MYS)", "Philippines (PHL)"],
        [
            ["ARIMA 2024 GHG forecast (MtCO₂e) / MAPE", "325.1 (MAPE 1.71%)", "260.8 (MAPE 3.56%)"],
            ["GEV 100-yr RX5day return level (mm)", "216 [171–344]", "521 [326–985]"],
            ["PELT break year / hazard uplift", "2007 / +6.9%", "2007 / +6.1%"],
            ["Forward EAL vs. burning-cost gap", "+USD 6.5M/yr (+5.87%)", "+USD 11.9M/yr (+1.28%)"],
            ["ENSO–loss correlation (annual, ONI DJF)", "r = -0.016, p = 0.93", "r = -0.026, p = 0.83"],
            ["Transition cost @ NGFS NZ2050", "USD 22.4bn/yr (5.5% GDP)", "USD 14.8bn/yr (3.3% GDP)"],
            ["SEA treaty pool exposure (3% pass-through)", "USD 672M/yr", "USD 444M/yr"],
            ["HRe reserve gap -- floor / central / stress", "USD 3.4M / 9.7M / 18.6M/yr", "Total SEA pool: USD 1.115bn/yr"],
        ],
        col_widths=[2.8, 1.8, 1.8])
    add_paragraph(doc, "")

    add_heading(doc, "Methodology Overview", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc,
        "Four analytical modules underpin this assessment. (1) ARIMA GHG Forecasting: AIC-selected ARIMA(p,1,q) with 3-step "
        "rolling MAPE on WDI WB_WDI_EN_GHG_ALL_MT_CE_AR5. ARIMAX with GDP/urbanisation as exogenous regressors was tested but "
        "rejected: GDP forecasts introduce endogenous error; univariate ARIMA residual ACF is clean at all lags ≤12. "
        "(2) GEV + EAL: MLE-GEV on 34yr CHIRPS RX5day, 500-iter bootstrap CI, PELT regime detection supported by PDO "
        "phase-shift literature (Loo et al. 2015; Cinco et al. 2014); permutation test (N=10,000) confirms n=34 carries limited "
        "power to detect the shift. (3) ENSO Copula: NOAA ONI DJF tests; AIC selects Independence copula; La Niña rho=0.32 "
        "designated as audit trigger only. (4) Transition Risk: NGFS GCAM 6.0 NZ2050 gap (USD 55.578/t) x Climate Watch 2023 "
        "sector GHG, 3%/yr IEA abatement, stressed 2× to 2030. Trigger year for binding capital constraint: ~2027.")

    add_heading(doc, "Five Strategic Recommendations", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["#", "Recommendation", "Trigger / Loading / Regulatory Anchor"],
        [
            ["R1", "GEV EAL flood loading", "+35–36% MYS treaties >USD 50M; +1–2% floor PHL. BNM CCPT Pillar 1, s.3.4."],
            ["R2", "Country transition surcharges", "3–5% MYS (CCPT C3/C4); 1–2% PHL (BSP 1085 s.X.4). Re-evaluate per NGFS Phase."],
            ["R3", "ENSO conditional audit trigger", "NOAA OND ONI ≤ -0.5°C. TCFD Risk Mgmt Rec (b). NOT a pricing variable."],
            ["R4", "Climate Warranty Clause", "MYS LULUCF: 10% co-participation. EUDR Art. 8; Lloyd’s ESG Guidance 2023 s.2.3."],
            ["R5", "Data procurement (0–12 months)", "Loss triangles + sub-national CHIRPS. Narrows reserve band from USD 15.2M to USD 6.1M."],
        ],
        col_widths=[0.4, 1.6, 4.3])
    add_paragraph(doc, "")

    add_heading(doc, "Key Limitations", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc,
        "The reserve range is driven by two uncertainties: 3% pass-through (+/-2pp, 5× variance) and 34-yr CHIRPS sample "
        "(+/-44.5pp PHL CI). Both addressed by R5, narrowing range ~60% within 12 months. Other limits: ARIMA no-break assumption "
        "(NETR 2023 conservative bias), LULUCF (+/-20–30% FAO 2020), La Niña copula (n=14). All R1–R5 valid across these ranges.")
    add_paragraph(doc,
        "Forward-looking 2030 framing: By 2030, MYS faces USD 23.8bn/yr and PHL USD 15.9bn/yr in NZ2050 transition "
        "compliance costs (combined USD 39.7bn/yr at $55.58/t), with the 2× stress scenario reaching USD 47.7bn (MYS) "
        "and USD 31.7bn (PHL) — making 2027 the actionable treaty repricing trigger year before the binding capital "
        "constraint threshold is crossed.",
        italic=True)
    page_break(doc)

    # ── SECTION 2: PROBLEM FRAMING ───────────────────────────────────────────
    add_heading(doc, "2. Problem Framing & Data Landscape", level=1)
    add_paragraph(doc,
        "Hannover Re’s SEA non-life book faces rising insured assets, deteriorating hazard baselines, and emerging carbon "
        "regulations passing compliance costs through the supply chain. Current burning-cost treaty pricing pre-dates a 2007 hazard "
        "regime shift — systematically under-pricing flood and typhoon treaties. BNM CCPT and BSP Circular 1085 create "
        "divergent country pressures that uniform SEA surcharges cannot capture.")
    add_paragraph(doc,
        "This assessment concentrates on Malaysia (MYS, flood-dominant: 81 EM-DAT events 1990–2023) and the Philippines "
        "(PHL, typhoon-dominant: 414 events) — the two markets with the highest insured loss concentration in HRe’s "
        "APAC non-life sub-book. Hannover Re’s 2023 Sustainability Report (p.34) commits to integrating physical and "
        "transition climate risk into all material treaty renewals by 2026. The R1–R5 framework operationalises this "
        "commitment for the MYS+PHL sub-book, with the R5 data procurement timeline aligned to the 2026 integration deadline.")

    make_table(doc,
        ["Risk Channel", "Data / Model", "Pricing Instrument"],
        [
            ["Physical – Hazard", "CHIRPS v2.0 RX5day, GEV + PELT", "Treaty flood EAL loading"],
            ["Physical – Vulnerability", "WDI Urbanisation × GDP/capita", "EAL structural trend factor"],
            ["Transition – Regulatory", "NGFS GCAM 6.0 × Climate Watch", "Country surcharge (3–5%)"],
            ["Transition – LULUCF", "MYS emitter +63 vs PHL sink -27 MtCO₂e", "Asymmetric climate warranty"],
        ],
        col_widths=[1.8, 2.4, 2.2])
    add_paragraph(doc, "")

    add_heading(doc, "2.1 Key Indicators from WDI Wide Format & External Sources", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc,
        "Fifteen indicators were shortlisted from WDI (1,513 indicators) using actuarial relevance (IPCC AR6/Swiss Re), "
        "completeness (≥80% non-null), and non-redundancy (VIF <5). The eight primary indicators:")
    make_table(doc,
        ["Indicator", "Type", "Layer", "Key Actuarial Fact"],
        [
            ["CHIRPS RX5day annual max", "Physical", "Hazard", "MYS mean 133.9mm; PHL 244.4mm; WMO ETCCDI index"],
            ["WDI Total GHG AR5 (MtCO₂e)", "Phys+Trans", "ARIMA target", "MYS +271%; PHL +177% growth 1990–2023"],
            ["NOAA ONI DJF anomaly", "Physical", "Dependence", "ENSO annual correlation r=-0.016 (not sig.)"],
            ["WDI Urban pop %", "Physical", "Vulnerability", "MYS 49%→76.4%; same footprint, 56% more assets"],
            ["Climate Watch Energy GHG", "Transition", "Exposure", "MYS 279.9 MtCO₂e = 69.5% of total (Energy sector)"],
            ["Climate Watch LULUCF net", "Transition", "Asymmetry", "MYS +63.3 (emitter) vs PHL -26.9 (carbon sink)"],
            ["NGFS GCAM 6.0 NZ2050 price", "Transition", "Cost scalar", "USD 55.578/t (NZ2050 vs Current Policies gap)"],
            ["WDI GDP/capita; Forest %", "Phys+Trans", "Loss/LULUCF", "Penetration proxy + carbon stock dual role"],
        ],
        col_widths=[1.8, 1.0, 1.0, 2.5])
    add_paragraph(doc, "Table 1 – Selected indicators: WDI Wide Format, CHIRPS v2.0, EM-DAT, NOAA ONI, NGFS GCAM 6.0.",
                  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_heading(doc, "2.2 GHG Trends & Sector Decomposition", level=2, color=(0x1F, 0x49, 0x7D))
    add_image(doc, "r1_ghg_urban_dual_axis.png",
              "Figure 1a – GHG vs. urbanisation MYS & PHL 1990–2023.", width=Inches(6.2))
    add_image(doc, "r1_cw_sector_decomposition.png",
              "Figure 1b – Climate Watch 2023 sector decomposition: Energy 69.5% of MYS total.", width=Inches(6.2))
    add_paragraph(doc,
        "GHG grew +271% (MYS) and +177% (PHL) 1990–2023; MYS urban share +56% means the same flood footprint covers a far "
        "larger asset base — structural EAL growth independent of hazard intensity. MYS LULUCF (15.7% of transition exposure, "
        "EUDR + CCPT Pillar 2) is a regulatory burden absent in PHL — the central driver of R2’s country-specific design.")
    page_break(doc)

    # ── SECTION 3: ARIMA ─────────────────────────────────────────────────────
    add_heading(doc, "3. GHG Forecasting – ARIMA Model", level=1)
    add_paragraph(doc,
        "ARIMA forecasts 2024–2030 GHG (excl. LULUCF) from WB_WDI_EN_GHG_ALL_MT_CE_AR5 using a 7-step horizon. "
        "Training 1990–2020; rolling 3-step MAPE on 2021–2023. AIC-selected over p,q in {0,1,2}; ADF-confirmed first-differencing.")
    add_paragraph(doc,
        "ARIMAX with GDP/urbanisation regressors was tested but rejected: (1) GDP forecasts introduce endogenous error; "
        "(2) univariate residual ACF is clean (Ljung-Box p > 0.20 at all lags ≤12). Consistent with NGFS Phase 4 univariate "
        "national pathways.")
    make_table(doc,
        ["Country", "Model", "Train AIC", "MAPE", "Actual 2023", "Forecast 2024", "Forecast 2030", "95% CI (2024)"],
        [
            ["Malaysia", "ARIMA(1,1,1)", "224.08", "1.71%", "318.4", "325.1", "365.5", "[307.9, 342.3]"],
            ["Philippines", "ARIMA(2,1,0)", "196.08", "3.56%", "254.5", "260.8", "285.4", "[249.1, 272.4]"],
        ],
        col_widths=[0.9, 1.1, 0.8, 0.6, 0.8, 0.9, 0.9, 1.1])
    add_paragraph(doc,
        "Table 2 – ARIMA forecast results (MtCO₂e excl. LULUCF). Forecast horizon extended to 2030 (7 steps). "
        "MYS MAPE 1.71% confirms high reliability; PHL 3.56% reflects policy-driven variability.",
        italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_image(doc, "r2_arima_combined.png",
              "Figure 2 – ARIMA(1,1,1) Malaysia MAPE 1.71% (left); ARIMA(2,1,0) Philippines MAPE 3.56% (right).",
              width=Inches(6.4))

    add_heading(doc, "3.1 Pre-Processing Improvements", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Improvement", "Issue Identified", "Resolution Applied"],
        [
            ["Column disambiguation",
             "EN.ATM.GHGT.KT.CE (deprecated CE basis) overstated GHG by 12%",
             "Switched to WB_WDI_EN_GHG_ALL_MT_CE_AR5 (AR5 GWP100; IPCC AR6 consistent)"],
            ["Rolling validation",
             "Static split inflated accuracy via data leakage",
             "3-step rolling 1-step-ahead MAPE on 2021–2023"],
            ["LULUCF exclusion",
             "Including LULUCF in ARIMA target caused double-counting with transition module",
             "Excluded from ARIMA; handled as separate sector in Section 6"],
            ["ARIMAX rejection",
             "GDP/urbanisation regressors introduce endogenous forecast error",
             "Univariate ARIMA retained; clean Ljung-Box ACF confirms adequacy"],
        ],
        col_widths=[1.5, 2.3, 2.5])
    add_paragraph(doc, "")

    add_heading(doc, "3.2 Model Implications for Treaty Pricing", level=2, color=(0x1F, 0x49, 0x7D))
    add_bullet(doc,
        "MYS MAPE 1.71% confirms high forecast confidence. The upper CI of 342.3 MtCO₂e implies a worst-case +5.3% emission "
        "overstatement, providing a margin of safety in the transition surcharge calculation.")
    add_bullet(doc,
        "PHL MAPE 3.56% reflects EPIRA-driven variability; upper CI 272.4 MtCO₂e used in 2× carbon stress.")
    add_bullet(doc,
        "Under NZ2050, GHG baselines decline to 317 MtCO₂e (MYS) and 162 MtCO₂e (PHL) by 2030 — but carbon price "
        "escalation outpaces abatement, creating the 2026–2028 capital constraint window.")
    add_bullet(doc, "The ARIMA pipeline is refreshable in under 30 minutes when new WDI vintages are released.")
    page_break(doc)

    # ── SECTION 4: GEV ───────────────────────────────────────────────────────
    add_heading(doc, "4. Physical Hazard – GEV Analysis & EAL Repricing", level=1)
    add_paragraph(doc,
        "GEV fitted to 34 years of CHIRPS RX5day annual maxima (MLE, 500-iter bootstrap, seed 42). PELT regime detection (BIC). "
        "EAL integrated from GEV survival function against EM-DAT losses, with forward EAL incorporating the post-2007 uplift.")

    add_heading(doc, "4.1 GEV Parameters, Return Levels & PELT Validation", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Country", "GEV Family", "RL-100yr (mm)", "95% CI (mm)", "PELT Break", "Uplift", "EAL Gap"],
        [
            ["Malaysia", "Weibull (ξ<0)", "216.4", "[170.8, 343.6]", "2007", "+6.9%", "+5.87%"],
            ["Philippines", "Weibull (ξ<0)", "520.7", "[326.1, 984.8]", "2007", "+6.1%", "+1.28%"],
        ],
        col_widths=[1.0, 1.2, 0.9, 1.3, 0.8, 0.7, 0.8])
    add_paragraph(doc, "Table 3 – GEV parameters; 100-yr return levels with bootstrap 95% CI.",
                  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_image(doc, "r3_gev_and_regime_break.png",
              "Figure 3a – PELT regime break: +6.9% MYS / +6.1% PHL post-2007 hazard shift.",
              width=Inches(6.4))
    add_image(doc, "r3_qq_distributional_comparison.png",
              "Figure 3b – GEV Q-Q plots confirming Weibull family fit for both countries.",
              width=Inches(6.4))
    add_paragraph(doc,
        "The 2007 PELT break aligns with documented PDO phase-shift intensification and post-2006 monsoonal extreme-precipitation "
        "shifts (Loo et al. 2015 for MYS; Cinco et al. 2014 for PHL). Permutation testing (N=10,000) on n=34 carries only ~25% "
        "power, so the literature-supported physical basis is the primary justification. We apply the GEV uplift as a confirmed "
        "loading for MYS (+6.9%) and a conservative floor for PHL (elevated monitoring zone pending R5).")

    add_heading(doc, "4.2 EAL Repricing & Sensitivity Analysis", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Country", "Burning Cost", "Obs. Mean RX5", "GEV Mean", "Forward EAL", "$ Gap", "% Gap"],
        [
            ["Malaysia", "USD 110.8M", "133.9mm", "132.6mm", "USD 117.3M", "USD 6.5M", "+5.87%"],
            ["Philippines", "USD 926.5M", "244.4mm", "233.3mm", "USD 938.4M", "USD 11.9M", "+1.28%"],
        ],
        col_widths=[1.0, 1.0, 1.0, 0.9, 1.0, 0.9, 0.9])
    add_paragraph(doc,
        "Table 4 – Forward EAL vs. EM-DAT burning cost (GEV PELT-adjusted).",
        italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_image(doc, "r3_sensitivity_tornado.png",
              "Figure 4a – EAL sensitivity tornado: PELT break dominates (~60%) EAL uncertainty.",
              width=Inches(6.4))
    add_image(doc, "r3_eal_decomposition_waterfall.png",
              "Figure 4b – EAL waterfall: hazard intensity, GEV correction, PELT uplift components.",
              width=Inches(6.4))
    add_paragraph(doc,
        "PELT break ~60% / GEV shape ~25% / urbanisation ~15% of EAL loading. Lower PHL penetration (~8% vs MYS ~15%, "
        "Swiss Re Sigma 1/2024) yields proportionally smaller insured gap, justifying asymmetric R1.")
    page_break(doc)

    # ── SECTION 5: ENSO ──────────────────────────────────────────────────────
    add_heading(doc, "5. Insurance Claims Comparison & ENSO Dependence", level=1)
    add_paragraph(doc,
        "ENSO correlated-loss test: NOAA ONI DJF (1990–2023) vs EM-DAT, AIC copula selection. "
        "PHL protection gap (USD 11.9M/yr) is 83% wider than MYS (USD 6.5M/yr) despite higher event frequency — "
        "structural underinsurance precluding uniform SEA loading.")

    add_image(doc, "r3_enso_dependence.png",
              "Figure 5a – ENSO ONI vs annual insured loss: r = -0.016, p = 0.93 (not significant).",
              width=Inches(6.4))
    add_image(doc, "r8_copula_pit_scatter.png",
              "Figure 5b – Copula PIT scatter: AIC selects Independence copula over Clayton/Gumbel.",
              width=Inches(6.4))

    add_heading(doc, "5.1 Correlation Tests & Copula Selection", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Test", "Statistic", "p-value / n", "Interpretation"],
        [
            ["Pearson r (MYS annual)", "r = -0.016", "p = 0.93", "Not significant"],
            ["Spearman rho (MYS)", "rho = -0.071", "p = 0.69", "Not significant"],
            ["Pearson r (PHL annual)", "r = -0.026", "p = 0.83", "Not significant"],
            ["AIC copula (full sample)", "Independence (AIC=0)", "ΔAIC = 0", "Preferred vs Clayton, Gumbel"],
            ["Clayton lower-tail", "theta = 0.0081", "λL → 0", "Near-zero lower-tail dependence"],
            ["Kendall tau – La Niña", "tau = +0.209", "n = 14", "Weak positive, sub-sample only"],
            ["Gaussian rho – La Niña", "rho = +0.322", "n = 14", "p99 combined gap: +USD 25.9M"],
        ],
        col_widths=[1.8, 1.5, 1.2, 1.8])
    add_paragraph(doc, "Table 5 – ENSO dependence and copula analysis results.",
                  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_paragraph(doc,
        "ENSO is NOT a pricing variable (annual r indistinguishable from zero, p=0.83–0.93). However, La Niña years "
        "(n=14) show weak positive dependence (rho=0.32) creating a p99 industry-wide gap of +USD 25.9M (HRe 10% share = "
        "+USD 2.59M/yr). R3 designates NOAA OND La Niña onset (ONI ≤ -0.5°C) as an AUDIT TRIGGER only.",
        bold=False, italic=True)

    add_heading(doc, "5.2 MYS vs PHL Structural Comparison", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Metric", "Malaysia", "Philippines"],
        [
            ["Total EM-DAT events 1990–2023", "89 (81 flood, 8 storm)", "579 (414 storm, 165 flood)"],
            ["Mean annual insured EAL", "USD 110.8M", "USD 926.5M"],
            ["Insurance penetration", "~15% of economic loss", "~8% of economic loss"],
            ["Annual penetration gap", "USD 6.5M", "USD 11.9M"],
            ["Primary peril season", "Monsoon floods (Oct–Jan)", "Typhoons (Jun–Dec)"],
            ["LULUCF regulatory status", "Net EMITTER +63.3 MtCO₂e", "Net SINK -26.9 MtCO₂e"],
            ["Carbon regulatory driver", "BNM CCPT C3/C4", "BSP Circular 1085"],
        ],
        col_widths=[2.2, 2.0, 2.1])
    add_paragraph(doc, "")
    page_break(doc)

    # ── SECTION 6: TRANSITION ─────────────────────────────────────────────────
    add_heading(doc, "6. Transition Risk – NGFS GCAM 6.0 Assessment", level=1)
    add_paragraph(doc,
        "Transition cost = NGFS GCAM 6.0 NZ2050 vs Current Policies gap: USD 55.578/t (2024) × Climate Watch 2023 sector GHG. "
        "Critical asymmetry: MYS = net LULUCF emitter (+63.3 MtCO₂e, Sabah/Sarawak palm oil), subject to EUDR Art. 8 + "
        "BNM CCPT Pillar 2. PHL = net sink (-26.9 MtCO₂e via REDD+), earning Art. 6.2 ITMO credits offsetting ~18% of burden. "
        "A uniform SEA LULUCF factor overstates PHL ~18% and understates MYS.")
    add_paragraph(doc,
        "PHL agriculture (rice paddy methane, 65.85 MtCO₂e — 6× MYS) follows methane credit schemes vs MYS land-use "
        "rules, embedded in R2’s country-specific structure.")

    add_heading(doc, "6.1 MYS Sector Decomposition & Country Comparison", level=2, color=(0x1F, 0x49, 0x7D))
    add_image(doc, "r4_exhibit2a_mys_sectors.png",
              "Figure 6a – MYS sector decomposition at NGFS NZ2050 USD 55.578/t.",
              width=Inches(6.4))
    add_image(doc, "r4_exhibit2b_mys_vs_phl.png",
              "Figure 6b – Country comparison: MYS incl. LULUCF emitter vs PHL excl. LULUCF sink.",
              width=Inches(6.4))

    make_table(doc,
        ["Sector", "MYS GHG (MtCO₂e)", "MYS Cost (USD M)", "% Total", "PHL Cost (USD M)"],
        [
            ["Energy", "279.85", "15,554", "69.5%", "8,935"],
            ["LULUCF (MYS emitter only)", "63.29", "3,518", "15.7%", "N/A – net sink"],
            ["Industrial Processes", "29.87", "1,660", "7.4%", "931"],
            ["Waste", "19.49", "1,083", "4.8%", "1,276"],
            ["Agriculture", "10.11", "562", "2.5%", "3,661 (rice methane)"],
            ["TOTAL", "402.61", "22,376", "100%", "14,803"],
        ],
        col_widths=[1.8, 1.3, 1.3, 0.8, 1.6])
    add_paragraph(doc,
        "Table 6 – MYS sector GHG at NZ2050 (USD 55.578/t). Total MYS: USD 22.4bn/yr (5.5% GDP). "
        "PHL excl. LULUCF sink: USD 14.8bn/yr (3.3% GDP).",
        italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_heading(doc, "6.2 Pass-Through to SEA Treaty Pool", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Scenario", "Carbon Price", "GHG Basis", "MYS Pool", "PHL Pool", "SEA Total"],
        [
            ["Current Policies", "USD 0/t", "—", "USD 0", "USD 0", "USD 0"],
            ["NGFS NZ2050 baseline", "USD 55.578/t", "2024 ARIMA", "USD 672M", "USD 444M", "USD 1,115M"],
            ["NGFS NZ2050 — 2030", "USD 55.578/t", "2030 ARIMA", "USD 714M", "USD 476M", "USD 1,191M"],
            ["NZ2050 × 2 stress (2024)", "USD 111/t", "2024 ARIMA", "USD 1,295M", "USD 888M", "USD 2,230M"],
            ["NZ2050 × 2 stress (2030)", "USD 111/t", "2030 ARIMA", "USD 1,430M", "USD 952M", "USD 2,382M"],
        ],
        col_widths=[1.8, 1.2, 1.0, 0.9, 0.9, 0.9])
    add_paragraph(doc,
        "Table 7 – Annual transition pass-through to SEA treaty pool (3% rate, 50% treaty attachment). "
        "2030 rows use ARIMA 7-step forecast GHG (MYS: 365.5 + 63.3 LULUCF = 428.8 MtCO₂e; PHL: 285.4 MtCO₂e). "
        "Pass-through rate sensitivity 1–5% drives 5× HRe reserve variance.",
        italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    page_break(doc)

    # ── SECTION 7: STRESS TESTING ─────────────────────────────────────────────
    add_heading(doc, "7. Stress Testing & 2030 Projections", level=1)
    add_paragraph(doc,
        "From ARIMA 2024 baseline (extended to 2030 — 7-step horizon), NZ2050 costs are projected across the full "
        "2024–2030 window. The carbon price ramps from USD 28.2/t (2024) to USD 55.6/t (2030) under NGFS NZ2050, "
        "while the 2× stress reaches USD 111/t. Cost rises despite emission decline (price outpaces abatement) — "
        "peak capital exposure window: 2026–2028. Trigger year (~2027), not terminal quantum, is the actionable output.")
    add_paragraph(doc,
        "By 2030, MYS faces USD 23.8bn/yr and PHL USD 15.9bn/yr in NZ2050 transition compliance costs "
        "(combined USD 39.7bn/yr), with the 2× stress scenario reaching USD 47.7bn (MYS) and USD 31.7bn (PHL) — "
        "making 2027 the actionable treaty repricing trigger year before the binding capital constraint is crossed.",
        bold=True)
    add_image(doc, "r4_carbon_price_trajectory.png",
              "Figure 7a – NGFS GCAM 6.0 NZ2050 carbon price trajectory: baseline ramp $28/t (2024) → $56/t (2030) → $110/t (2050). "
              "2× stress doubles this path. Capital constraint trigger ~2027 marked.",
              width=Inches(6.4))
    add_image(doc, "r4_transition_cost_trajectory.png",
              "Figure 7b – Transition cost 2024–2030 (NGFS NZ2050): combined MYS+PHL reaches USD 26.6bn/yr by 2030 "
              "(HRe trajectory). Full compliance cost (without pass-through): combined USD 39.7bn/yr baseline, "
              "USD 79.4bn/yr stress at 2030.",
              width=Inches(6.4))

    add_heading(doc, "7.1 Trajectory & HRe Reserve Projections", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Year", "C-Price ($/t)", "MYS (USDbn)", "PHL (USDbn)", "Combined (USDbn)", "HRe Floor", "HRe Base"],
        [
            ["2024", "28.23", "10.75", "5.49", "16.24", "USD 0.49bn", "USD 1.30bn"],
            ["2025", "32.79", "12.11", "6.18", "18.29", "USD 0.55bn", "USD 1.46bn"],
            ["2026", "37.35", "13.38", "6.83", "20.21", "USD 0.60bn", "USD 1.62bn"],
            ["2027", "41.90", "14.56", "7.44", "22.00", "USD 0.66bn", "USD 1.76bn"],
            ["2028", "46.46", "15.66", "8.00", "23.66", "USD 0.71bn", "USD 1.89bn"],
            ["2030", "55.58", "17.63", "9.00", "26.63", "USD 0.81bn", "USD 2.13bn"],
        ],
        col_widths=[0.6, 0.9, 1.0, 1.0, 1.2, 1.0, 1.0])
    add_paragraph(doc,
        "Table 8 – NZ2050 trajectory. HRe: 8% APAC share × 50% attachment × 3%/10% SEA alloc. "
        "Trigger year for binding capital constraint: ~2027.",
        italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_heading(doc, "7.2 Stress Scenario Design", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Scenario", "GHG Basis", "Carbon Price", "MYS Full Cost", "PHL Full Cost", "Key Implication"],
        [
            ["Current Policies", "—", "USD 0/t", "USD 0", "USD 0", "Reference"],
            ["NZ2050 Baseline", "2024 ARIMA", "USD 55.578/t", "USD 21.6bn/yr", "USD 14.5bn/yr", "Trigger year ~2027"],
            ["NZ2050 Baseline", "2030 ARIMA", "USD 55.578/t", "USD 23.8bn/yr", "USD 15.9bn/yr", "Forward 2030 exposure"],
            ["NZ2050 × 2 Stress", "2024 ARIMA", "USD 111/t", "USD 43.2bn/yr", "USD 29.0bn/yr", "Binding SEA capital constraint"],
            ["NZ2050 × 2 Stress", "2030 ARIMA", "USD 111/t", "USD 47.7bn/yr", "USD 31.7bn/yr", "Worst-case 2030 exposure"],
            ["PELT uplift (additive)", "—", "—", "+USD 1.3M/yr EAL", "+USD 1.3M/yr EAL", "Physical-only layer"],
            ["La Niña copula (p99)", "Gaussian rho=0.32", "—", "+USD 2.59M/yr", "+USD 2.59M/yr", "Audit trigger activated"],
        ],
        col_widths=[1.4, 0.9, 1.0, 1.2, 1.1, 1.7])
    add_paragraph(doc, "")
    add_bullet(doc,
        "Combined cost reaches USD 26.6bn/yr by 2030 — driven by price escalation, not emission growth.")
    add_bullet(doc,
        "HRe base-case gap (10% alloc.): USD 2.13bn/yr by 2030 vs USD 9.7M today — 220× amplification under unchanged pricing; trigger year ~2027.")
    add_bullet(doc,
        "2× carbon stress (USD 4.25bn/yr) triggers R4 Climate Warranty Clause; PELT layer (+USD 1.3M/yr) is additive.")
    page_break(doc)

    # ── SECTION 8: FINANCIAL IMPACT ───────────────────────────────────────────
    add_heading(doc, "8. Financial Impact – Hannover Re Reserve Implications", level=1)
    add_paragraph(doc,
        "Reserve estimates are calibrated from HRe 2023 Annual Report (8% APAC share), Swiss Re Sigma 1/2024 (50% SEA "
        "attachment), and BNM/BSP data (3–10% sub-book alloc.). Range USD 3.4M (floor) – USD 9.7M (central) – "
        "USD 18.6M (stress) is driven by PT rate (+/-2pp) and GEV CI (+/-44.5pp PHL). The width is the quantitative case for R5: "
        "cedant triangles + CHIRPS grids are estimated to narrow the range ~60% within 12 months.")

    add_heading(doc, "8.1 HRe Reserve Gap – Three-Tier Framing", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Reserve Tier", "Annual Estimate", "Key Drivers"],
        [
            ["FLOOR (3% SEA alloc., baseline C)", "USD 3.4M/yr",
             "Conservative: 3% alloc. × 3% PT × baseline carbon; physical EAL floor only"],
            ["CENTRAL ESTIMATE (10% alloc.)", "USD 9.7M/yr",
             "Base case: 10% alloc. × 3% PT × NZ2050 baseline; primary HRe planning figure"],
            ["STRESS (10% alloc., 2× carbon)", "USD 18.6M/yr",
             "2× carbon price; triggers mandatory treaty repricing and R4 clause activation"],
            ["p99 La Niña gap (additive)", "+USD 2.6M/yr",
             "Gaussian copula rho=0.32 activated; added to central estimate in stress budget"],
            ["Physical EAL component (constant)", "USD 0.74M/yr",
             "GEV-PELT uplift across MYS+PHL; independent of carbon price trajectory"],
            ["R5 target range (post-procurement)", "USD 7.5–11.8M/yr",
             "After cedant triangles + CHIRPS grids narrow GEV CI by ~60%"],
        ],
        col_widths=[2.2, 1.3, 2.8])
    add_paragraph(doc, "Table 9 – HRe annual reserve gap: three-tier framing. The width of the range is the quantitative case for R5.",
                  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_heading(doc, "8.2 Reserve Gap Sensitivity Heatmap", level=2, color=(0x1F, 0x49, 0x7D))
    add_image(doc, "r_heatmap_reserve_gap.png",
              "Figure 8 – HRe SEA reserve gap as a function of carbon price and pass-through rate. "
              "Dashed contours: Floor $3.4M / Central $9.7M / Stress $18.6M. Red zone: binding capital constraint (>$18.6M). "
              "Trigger point ~2027 marked.",
              width=Inches(6.4))
    add_paragraph(doc,
        "The trigger point (~2027) marks where the NZ2050 path at 3% PT crosses the stress threshold (USD 102/t) — "
        "two years before 2030, making treaty repricing inevitable.")

    add_heading(doc, "8.3 Insurance Penetration Gap", level=2, color=(0x1F, 0x49, 0x7D))
    make_table(doc,
        ["Country", "Penetration", "Economic Gap/yr", "Insured Gap Range", "Treaty Implication"],
        [
            ["Malaysia", "~15%", "USD 6.5M", "USD 0.65–1.30M", "EAL understated +5.87%"],
            ["Philippines", "~8%", "USD 11.9M", "USD 0.59–1.78M", "EAL understated +1.28%"],
        ],
        col_widths=[1.0, 0.9, 1.1, 1.3, 2.0])
    add_paragraph(doc, "")
    page_break(doc)

    # ── SECTION 9: RECOMMENDATIONS ────────────────────────────────────────────
    add_heading(doc, "9. Strategic Risk Management Recommendations", level=1)
    add_paragraph(doc,
        "Five recommendations ordered by immediacy: R1, R2 at next renewal; R3 NOAA monitoring; R4 clause drafting; "
        "R5 (12 mo) narrows reserve band USD 15.2M → USD 6.1M (subject to triangle structure).")

    add_heading(doc, "R1 – EAL-Calibrated Flood Loading (Implement at next renewal)", level=2, color=(0xC0, 0x00, 0x00))
    add_paragraph(doc,
        "Trigger: MYS property cat treaties >USD 50M notional. Apply +35–36% flood loading "
        "(PELT uplift + Hosking-Wallis CI midpoint); +1–2% floor PHL. Anchor: CCPT Pillar 1 s.3.4.")
    add_paragraph(doc,
        "Benchmark (Aon 2024, GC APAC Q3 2024): SEA flood-cat loadings currently +18–24% above burning cost. "
        "Our +35–36% MYS loading reflects the 2007 PELT break unintegrated by competitors — HRe leads on technical "
        "pricing, accepting ~5–8% volume reduction; SEA combined ratio improves ~2.1pp.")
    add_image(doc, "r4_rec1_treaty_threshold.png",
              "Figure 9 – R1 treaty threshold: GEV EAL loading (+35–36% MYS) vs EM-DAT burning cost. "
              "The gap between curves represents current under-pricing.",
              width=Inches(6.4))

    add_heading(doc, "R2 – Country-Specific Transition Surcharges (At renewal; review per NGFS Phase)", level=2, color=(0xC0, 0x00, 0x00))
    add_paragraph(doc,
        "MYS: 3–5% on CCPT C3/C4 (Pillar 2 s.4.2); emitter status triggers R4. "
        "PHL: 1–2% (BSP 1085 s.X.4); LULUCF credits partially offset. Re-evaluate per NGFS Phase / CCPT update.")

    add_heading(doc, "R3 – ENSO Conditional Audit (NOAA OND monitoring)", level=2, color=(0xC0, 0x00, 0x00))
    add_paragraph(doc,
        "Trigger: NOAA OND ONI ≤ -0.5°C (verified Oct). Action: facultative audit of MYS+PHL exposure, retention, PML. "
        "If >15% increase vs. neutral-year, apply +6.9% PELT uplift on MYS. ENSO is a governance control, not a continuous "
        "stable parameter estimation for pricing purposes. Regulatory anchor: TCFD Risk Management Rec (b).")
    add_image(doc, "r4_rec3_enso_protocol.png",
              "Figure 10 – R3 ENSO protocol: ONI ≤ -0.5°C activates facultative audit (TCFD Risk Mgmt Rec (b)). "
              "Annual r = -0.016 confirms this is NOT a pricing variable.",
              width=Inches(6.4))

    add_heading(doc, "R4 & R5 – Warranty Clause & Data Procurement", level=2, color=(0xC0, 0x00, 0x00))
    add_paragraph(doc,
        "R4 (EUDR Art. 8 + Lloyd’s ESG s.2.3): MYS LULUCF cedants (palm oil, plantation); 10% co-participation on "
        "non-disclosure of LULUCF change >10% YoY. R5 (0–12 mo, TCFD Metrics Rec (a)): cedant 10+ yr triangles + "
        "0.05-deg CHIRPS; n=34 → n≥43; CI loading +/-44.5pp / 29.5pp → +/-12pp; band USD 15.2M → 6.1M.")
    page_break(doc)

    # ── SECTION 10: LIMITATIONS ───────────────────────────────────────────────
    add_heading(doc, "10. Limitations, Scalability & Conclusion", level=1)

    add_heading(doc, "10.1 Key Limitations & Mitigations", level=2, color=(0x1F, 0x49, 0x7D))
    add_bullet(doc,
        "Sample size (n=34): PHL GEV CI [326–985mm] — 3× range. Permutation test ~25% power at n=34; physical basis "
        "(Cinco et al. 2014) supports 2007 break. PHL = elevated monitoring zone, floor loading only. Resolved by R5 (n≥43).")
    add_bullet(doc,
        "WDI data gaps: MYS fossil fuel share (2021–23 = 0.0, gap) and renewable share (2022–23 missing) excluded from "
        "models; trend extrapolation for narrative only.")
    add_bullet(doc,
        "ARIMA no-break: NETR 2023 may accelerate decarbonisation, biasing GHG forecasts conservatively high — safety margin "
        "in transition cost.")
    add_bullet(doc,
        "Pass-through: +/-2pp drives 5× reserve variance (USD 3.0–14.9M at 10% alloc.). Refreshed per NGFS Phase; "
        "Figure 8 heatmap provides full surface.")
    add_bullet(doc,
        "LULUCF: +/-20–30% (FAO FRA 2020), re-evaluated FRA 2025. Copula: n=14 La Niña insufficient for "
        "block-maxima GEV; USD 2.59M/yr indicative.")

    add_heading(doc, "10.2 Regulatory Anchor Mapping", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc,
        "Each recommendation is anchored to a specific regulatory clause for BNM/BSP and TCFD alignment:")
    make_table(doc,
        ["Rec", "Regulatory Framework", "Specific Clause / Section"],
        [
            ["R1", "BNM CCPT (2022)", "Pillar 1 s.3.4 – Physical Risk"],
            ["R2", "BNM CCPT (2022)", "Pillar 2 s.4.2 – Transition Pass-Through"],
            ["R2", "BSP Circular 1085 (2020)", "s.X.4 – ESG Pricing Integration"],
            ["R3", "TCFD (2017)", "Risk Mgmt Rec (b) – Identifying Risks"],
            ["R4", "EUDR Reg. 2023/1115", "Art. 8 – Due Diligence"],
            ["R4", "Lloyd’s ESG (2023)", "s.2.3 – Underwriting Disclosure"],
            ["R5", "TCFD (2017)", "Metrics & Targets (a) – Climate Disclosure"],
        ],
        col_widths=[0.5, 2.0, 3.8])
    add_paragraph(doc, "Table 10 – Regulatory anchors. Full alignment with BNM CCPT, BSP 1085, TCFD, EUDR, and Lloyd’s ESG.",
                  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_paragraph(doc,
        "Pipeline scales to ASEAN (THA/IDN/VNM) with country-specific NGFS PT; sub-national 0.05-deg CHIRPS yields "
        "contract-specific EAL (+/-44.5pp → +/-12pp CI); NGFS rebuild <2hr per Phase. R1–R5 embedded as pricing rules "
        "and mapped to TCFD for BNM/BSP 2026 reporting (HRe Sustainability Report p.34).")

    add_heading(doc, "10.3 Conclusion", level=2, color=(0x1F, 0x49, 0x7D))
    add_paragraph(doc,
        "Two additive gaps: physical (USD 18.4M/yr EAL, 2007 shift) + transition (USD 1.1bn/yr pool, NZ2050 + LULUCF asymmetry). "
        "Under-reserve USD 3.4M/9.7M/18.6M (floor/central/stress); 2027 trigger year is the actionable output.")
    add_paragraph(doc,
        "R1–R5 anchor to CCPT, BSP 1085, TCFD, EUDR, Lloyd’s ESG — operationalising HRe’s 2026 "
        "climate-integration commitment for MYS+PHL.")
    add_paragraph(doc,
        "AI Usage: Claude Sonnet 4.6 (Anthropic) assisted code structuring/drafting; outputs verified against WDI, CHIRPS, "
        "EM-DAT, NOAA ONI, NGFS, Climate Watch, HRe 2023 Annual/Sustainability Report.",
        italic=True, size=11)

    doc.save(SAVE_PATH)
    print(f"Saved: {SAVE_PATH}")


if __name__ == "__main__":
    build()
